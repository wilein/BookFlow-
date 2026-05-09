# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT_PATH = Path("thesis_output") / "BookFlow_core_interface_db_tables.docx"


INTERFACE_ROWS = [
    ("用户认证", "POST", "/user/auth/wechat", "微信小程序登录，根据微信临时登录凭证换取用户身份并生成系统登录态。", "code、用户基本信息；返回 token、用户信息"),
    ("用户认证", "POST", "/user/auth/dev-login", "开发环境快速登录，用于本地调试接口鉴权流程。", "测试用户标识；返回 token"),
    ("用户认证", "GET", "/user/auth/check", "校验当前登录态是否有效。", "请求头携带 token"),
    ("用户认证", "POST", "/user/auth/logout", "退出登录，清理当前用户登录状态。", "请求头携带 token"),
    ("用户信息", "GET", "/user/info", "获取当前登录用户的基础信息。", "用户 id 从登录态中读取"),
    ("个人中心", "GET", "/user/profile", "查询用户资料与学生认证信息。", "返回昵称、头像、认证状态、信用分等"),
    ("个人中心", "GET", "/user/stats", "统计个人中心展示数据。", "返回书架、订单、收藏、批注、路径等数量"),
    ("个人中心", "POST", "/user/profile/update", "修改用户个人资料。", "昵称、头像、简介、手机号等"),
    ("学生认证", "POST", "/user/profile/verify-student", "提交学生认证资料，用于平台身份可信度管理。", "学号、真实姓名、学校、院系、学生证图片"),
    ("学生认证", "POST", "/user/profile/upload-image", "上传学生认证图片或个人资料图片。", "multipart 文件上传，返回图片地址"),
    ("地址管理", "GET", "/user/address/list", "查询当前用户收货地址列表。", "返回地址集合"),
    ("地址管理", "POST", "/user/address/save", "新增或修改收货地址。", "收件人、手机号、省市区、详细地址、默认标识"),
    ("地址管理", "POST", "/user/address/delete", "删除指定收货地址。", "addressId"),
    ("地址管理", "POST", "/user/address/set-default", "设置默认收货地址。", "addressId"),
    ("书籍管理", "GET", "/book/list", "分页查询二手书列表，支持首页展示和分类筛选。", "page、size、category、status 等"),
    ("书籍管理", "GET", "/book/category", "查询书籍分类数据。", "返回分类列表"),
    ("书籍管理", "GET", "/book/search", "按关键词搜索书籍。", "keyword、page、size"),
    ("书籍管理", "GET", "/book/detail", "查询书籍详情，包含卖家、书籍状态和统计信息。", "bookId"),
    ("书籍管理", "POST", "/book/publish", "发布二手书信息。", "ISBN、书名、作者、出版社、价格、成色、图片等"),
    ("书籍管理", "POST", "/book/update", "修改已发布书籍信息。", "bookId 及待更新字段"),
    ("书籍管理", "POST", "/book/change-status", "变更书籍状态，如上架、下架或售出。", "bookId、status"),
    ("书籍管理", "POST", "/book/upload-image", "上传书籍封面或详情图片。", "multipart 文件上传，返回图片地址"),
    ("订单交易", "POST", "/order/create", "买家基于书籍创建交易订单。", "bookId、收货地址、买家留言"),
    ("订单交易", "POST", "/order/cancel", "取消未完成订单。", "orderId、取消原因"),
    ("订单交易", "POST", "/order/pay/mock", "模拟支付订单，用于毕业设计演示交易流程。", "orderId"),
    ("订单交易", "POST", "/order/pay", "处理订单支付请求并更新支付状态。", "orderId、paymentMethod"),
    ("订单交易", "POST", "/order/ship", "卖家确认发货，推进订单状态。", "orderId"),
    ("订单交易", "POST", "/order/confirm-receipt", "买家确认收货，完成订单闭环。", "orderId"),
    ("订单交易", "GET", "/order/detail", "查询订单详情。", "orderId"),
    ("订单纠纷", "POST", "/order/issue/create", "提交订单问题或售后纠纷。", "orderId、type、content"),
    ("订单纠纷", "GET", "/order/issue/list", "查询订单相关问题记录。", "orderId"),
    ("订单纠纷", "POST", "/order/issue/reply", "回复订单问题并更新处理状态。", "issueId、replyContent、status"),
    ("书籍批注", "GET", "/annotation/list", "查询指定书籍的公开或个人批注。", "bookId、pageNum、visibility"),
    ("书籍批注", "POST", "/annotation/create", "创建书籍批注，支持文字和图片内容。", "bookId、pageNum、content、positionText、type、visibility"),
    ("书籍批注", "POST", "/annotation/upload-image", "上传批注图片。", "multipart 文件上传，返回图片地址"),
    ("书籍批注", "POST", "/annotation/toggle-like", "对批注进行点赞或取消点赞。", "annotationId"),
    ("资源管理", "GET", "/resource/my-list", "查询当前用户上传的学习资源。", "page、size"),
    ("资源管理", "GET", "/resource/list", "查询与书籍或路径绑定的资源列表。", "bookId、bindType、bindId"),
    ("资源管理", "POST", "/resource/upload-file", "上传 PDF、PPT、图片等学习资源文件。", "multipart 文件上传，返回文件地址和大小"),
    ("资源管理", "POST", "/resource/create", "新增资源元数据并建立资源归属关系。", "title、type、fileUrl、bookId、bindType、bindId"),
    ("资源管理", "POST", "/resource/update", "修改资源标题、描述、可见性等信息。", "resourceId 及待更新字段"),
    ("资源管理", "POST", "/resource/delete", "删除资源记录。", "resourceId"),
    ("学习路径", "GET", "/path/detail", "查询学习路径详情及节点树。", "pathId"),
    ("学习路径", "POST", "/path/save-draft", "保存学习路径草稿。", "title、description、bookId、nodes"),
    ("学习路径", "POST", "/path/publish", "发布学习路径，使其可被其他用户查看或复制。", "pathId 或完整路径数据"),
    ("学习路径", "POST", "/path/copy", "复制他人学习路径形成个人学习路线。", "sourcePathId"),
    ("学习路径", "POST", "/path/change-status", "调整学习路径状态。", "pathId、status"),
    ("学习进度", "POST", "/path/progress/start", "开始学习某条路径并创建进度记录。", "pathId"),
    ("学习进度", "POST", "/path/progress/complete-node", "标记路径节点完成并更新总体进度。", "pathId、nodeId"),
    ("社区模块", "GET", "/community/feed", "查询社区动态信息流。", "page、size、type"),
    ("社区模块", "GET", "/community/activity", "查询社区活跃数据。", "用户活跃、热门内容等"),
    ("社区模块", "POST", "/community/post/create", "发布社区帖子或学习经验分享。", "title、content、type、sharedPathId"),
    ("社区模块", "POST", "/community/post/toggle-like", "对帖子点赞或取消点赞。", "postId"),
    ("社区模块", "POST", "/community/post/toggle-favorite", "收藏或取消收藏社区帖子。", "postId"),
    ("社区评论", "GET", "/community/post/comment/list", "查询帖子评论列表。", "postId、page、size"),
    ("社区评论", "POST", "/community/post/comment/create", "发布帖子评论。", "postId、content"),
    ("内容治理", "POST", "/community/post/report", "举报社区帖子。", "postId、reason、content"),
    ("内容治理", "GET", "/community/my-reports", "查询当前用户提交的举报记录。", "page、size"),
    ("收藏模块", "POST", "/favorite/toggle", "收藏或取消收藏指定对象。", "targetType、targetId"),
    ("收藏模块", "GET", "/favorite/status", "查询指定对象是否已被当前用户收藏。", "targetType、targetId"),
    ("私信模块", "GET", "/chat/session/list", "查询当前用户私信会话列表。", "page、size"),
    ("私信模块", "POST", "/chat/session/open", "打开或创建买卖双方围绕某本书的会话。", "bookId、sellerId 或 buyerId"),
    ("私信模块", "GET", "/chat/message/list", "查询会话消息明细。", "sessionId、page、size"),
    ("私信模块", "POST", "/chat/message/send", "发送私信消息并更新会话摘要。", "sessionId、content"),
    ("公共数据", "GET", "/common/banner/list", "查询首页轮播图数据。", "返回启用状态的 banner 列表"),
    ("浏览记录", "GET", "/user/history/list", "查询用户浏览历史。", "page、size"),
    ("浏览记录", "POST", "/user/history/record", "记录用户浏览书籍行为。", "bookId"),
    ("浏览记录", "POST", "/user/history/delete", "删除单条浏览记录。", "historyId"),
    ("浏览记录", "POST", "/user/history/clear", "清空当前用户浏览记录。", "无"),
    ("反馈通知", "POST", "/user/feedback/submit", "提交用户反馈。", "type、content、contact"),
    ("反馈通知", "GET", "/user/notifications", "查询系统通知列表。", "page、size"),
    ("反馈通知", "POST", "/user/notification/read", "将指定通知标记为已读。", "notificationId"),
]


DB_ROWS = [
    ("wx_user", "用户基础表", "存储微信登录用户的基础身份、昵称、头像、手机号和登录时间。", "PK：id", "openid、session_key、nickname、avatar_url、mobile、last_login_time", "与 user_profile、user_address、book、post、chat_message 等表通过 user_id 关联。"),
    ("user_profile", "学生认证表", "存储学生认证资料、认证状态、信用分和个人简介。", "PK：id；FK：user_id", "student_id、real_name、school、department、auth_status、credit_score", "与 wx_user 一对一，支撑实名认证与可信交易。"),
    ("user_address", "收货地址表", "保存用户订单收货地址，支持默认地址设置。", "PK：id；FK：user_id", "receiver_name、receiver_phone、province、city、district、detail_address、is_default", "一个用户可维护多个地址。"),
    ("book", "书籍信息表", "记录二手书发布信息，包括 ISBN、书名、作者、价格、成色、状态及统计数据。", "PK：id；FK：user_id", "isbn、title、author、publisher、price、condition、status、annotation_count", "与发布用户、订单、批注、资源、学习路径关联，是平台核心业务表。"),
    ("order", "订单表", "记录买卖双方围绕书籍形成的交易订单及支付、发货、收货状态。", "PK：id；FK：book_id、buyer_id、seller_id", "order_no、total_amount、status、payment_time、delivery_time、receive_time", "关联 book、买家用户、卖家用户和 order_issue。"),
    ("order_issue", "订单问题表", "记录订单纠纷、售后问题和处理回复。", "PK：id；FK：order_id、user_id、reply_user_id", "type、content、reply_content、status、reply_time", "一个订单可对应多条问题记录。"),
    ("annotation", "书籍批注表", "记录用户对书籍页码或位置的文字、图片批注，支持公开或私有。", "PK：id；FK：book_id、user_id", "page_num、content、position_text、image_url、visibility、like_count", "关联 book 和 wx_user，并被 annotation_like 引用。"),
    ("annotation_like", "批注点赞表", "记录用户对批注的点赞关系，避免重复点赞。", "PK：id；FK：annotation_id、user_id", "annotation_id、user_id、create_time", "annotation 与 wx_user 的关联表。"),
    ("resource", "学习资源表", "保存用户上传的 PDF、PPT、图片等资料，可绑定书籍或学习路径节点。", "PK：id；FK：user_id、book_id", "bind_type、bind_id、title、type、file_url、file_size、visibility", "关联上传用户、书籍和绑定对象。"),
    ("learning_path", "学习路径表", "记录用户创建或复制的课程/书籍学习路线。", "PK：id；FK：user_id、book_id、source_path_id", "title、description、difficulty、estimated_hours、status", "与 path_node、user_path_progress 构成学习路径模块。"),
    ("path_node", "路径节点表", "记录学习路径下的层级节点、顺序和绑定资源。", "PK：id；FK：path_id、parent_id", "title、description、order_num、estimated_minutes、resource_ids", "一个 learning_path 下有多个 path_node，支持父子节点结构。"),
    ("user_path_progress", "学习进度表", "记录用户对某条学习路径的学习进度。", "PK：id；FK：user_id、path_id", "progress_percent、completed_count、started_at、last_learn_time", "关联用户和学习路径，用于个人学习进度展示。"),
    ("post", "社区帖子表", "记录社区动态、经验分享和学习路径分享帖。", "PK：id；FK：user_id、shared_path_id", "title、content、type、view_count、like_count、comment_count", "关联发布用户和可选的 learning_path。"),
    ("comment", "评论表", "记录社区帖子的用户评论。", "PK：id；FK：post_id、user_id", "content、like_count、create_time", "一个 post 可对应多条 comment。"),
    ("chat_session", "私信会话表", "记录买卖双方围绕某本书建立的私信会话。", "PK：id；FK：book_id、buyer_id、seller_id", "last_message、last_message_time、unread_count", "关联 book、买家用户、卖家用户和 chat_message。"),
    ("chat_message", "私信消息表", "记录私信会话中的消息明细和已读状态。", "PK：id；FK：session_id、sender_id", "content、is_read、create_time", "一个 chat_session 下有多条 chat_message。"),
]


FIELD_ROWS = [
    ("wx_user", "id、openid、session_key、nickname、avatar_url、mobile、last_login_time", "分别表示用户主键、微信身份标识、微信会话密钥、昵称、头像、手机号和最近登录时间。"),
    ("user_profile", "id、user_id、student_id、real_name、school、department、auth_status、credit_score", "分别表示认证记录主键、所属用户、学号、真实姓名、学校、院系、认证状态和信用分。"),
    ("user_address", "id、user_id、receiver_name、receiver_phone、province、city、district、detail_address、is_default", "分别表示地址主键、所属用户、收件人、手机号、省、市、区、详细地址和默认地址标记。"),
    ("book", "id、user_id、isbn、title、author、publisher、price、condition、status、annotation_count", "分别表示书籍主键、发布者、ISBN、书名、作者、出版社、价格、成色、状态和批注数量。"),
    ("order", "id、order_no、book_id、buyer_id、seller_id、total_amount、status、payment_time、delivery_time、receive_time", "分别表示订单主键、订单号、书籍、买家、卖家、订单金额、订单状态、支付时间、发货时间和收货时间。"),
    ("order_issue", "id、order_id、user_id、type、content、reply_content、reply_user_id、status", "分别表示问题主键、所属订单、提交用户、问题类型、问题内容、回复内容、回复用户和处理状态。"),
    ("annotation", "id、book_id、user_id、page_num、content、position_text、image_url、visibility、like_count", "分别表示批注主键、所属书籍、创建用户、页码、批注内容、定位文字、图片地址、可见性和点赞数。"),
    ("annotation_like", "id、annotation_id、user_id、create_time", "分别表示点赞记录主键、批注编号、点赞用户和创建时间。"),
    ("resource", "id、user_id、book_id、bind_type、bind_id、title、type、file_url、file_size、visibility", "分别表示资源主键、上传用户、关联书籍、绑定类型、绑定对象、标题、资源类型、文件地址、文件大小和可见性。"),
    ("learning_path", "id、user_id、book_id、source_path_id、title、description、difficulty、estimated_hours、status", "分别表示路径主键、创建用户、关联书籍、来源路径、标题、描述、难度、预计学习时长和状态。"),
    ("path_node", "id、path_id、parent_id、title、description、order_num、estimated_minutes、resource_ids", "分别表示节点主键、所属路径、父节点、标题、描述、排序、预计学习分钟数和绑定资源。"),
    ("user_path_progress", "id、user_id、path_id、progress_percent、completed_count、started_at、last_learn_time", "分别表示进度主键、学习用户、学习路径、完成百分比、已完成节点数、开始时间和最近学习时间。"),
    ("post", "id、user_id、title、content、type、shared_path_id、view_count、like_count、comment_count", "分别表示帖子主键、发布用户、标题、内容、类型、分享路径、浏览量、点赞数和评论数。"),
    ("comment", "id、post_id、user_id、content、like_count、create_time", "分别表示评论主键、所属帖子、评论用户、评论内容、点赞数和创建时间。"),
    ("chat_session", "id、book_id、buyer_id、seller_id、last_message、last_message_time、unread_count", "分别表示会话主键、关联书籍、买家、卖家、最近消息、最近消息时间和未读数。"),
    ("chat_message", "id、session_id、sender_id、content、is_read、create_time", "分别表示消息主键、所属会话、发送者、消息内容、是否已读和发送时间。"),
]


def apply_font(run, name="宋体", size=9, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_text(cell, text, size=8.5, bold=False):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(str(text))
    apply_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def caption(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    apply_font(run, name="黑体", size=10.5, bold=True)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    repeat_header(table.rows[0])
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, size=9, bold=True)
        shade_cell(table.rows[0].cells[idx], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Cm(width)
    add_borders(table)
    doc.add_paragraph()


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BookFlow 核心接口与数据库关键表说明")
    apply_font(run, name="黑体", size=16, bold=True)

    desc = doc.add_paragraph()
    desc.paragraph_format.first_line_indent = Cm(0.74)
    desc.paragraph_format.line_spacing = 1.5
    run = desc.add_run(
        "说明：以下内容根据 BookFlow 后端控制器接口与实体表结构整理，可用于本科毕业论文“系统详细设计”“接口设计”和“数据库设计”章节。"
        "接口表按模块归类展示，未单独列出控制器名称；数据库表选取系统核心业务表，并保留论文说明所需的关键字段。"
    )
    apply_font(run, size=10.5)

    caption(doc, "表4-1 核心接口说明表")
    add_table(doc, ["功能模块", "请求方式", "接口地址", "接口用途", "主要参数/说明"], INTERFACE_ROWS, [2.1, 1.4, 3.6, 5.6, 4.8])

    caption(doc, "表4-2 数据库关键表说明表")
    add_table(doc, ["表名", "中文名称", "表作用", "主键/外键", "关键字段", "主要关联关系"], DB_ROWS, [2.7, 2.0, 4.2, 2.9, 4.3, 3.7])

    caption(doc, "表4-3 数据库核心字段说明表")
    add_table(doc, ["表名", "核心字段", "字段含义"], FIELD_ROWS, [3.0, 7.3, 7.2])

    OUT_PATH.parent.mkdir(exist_ok=True)
    doc.save(OUT_PATH)
    print(str(OUT_PATH.resolve()))


if __name__ == "__main__":
    main()
