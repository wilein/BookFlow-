# -*- coding: utf-8 -*-
from pathlib import Path
import importlib.util
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "thesis_output"
DIAGRAM_DIR = OUT_DIR / "diagrams_rewritten"
OUT_DOC = OUT_DIR / "校园学术资源传承平台论文_重写版.docx"


def load_table_source():
    script = OUT_DIR / "generate_interface_db_doc.py"
    spec = importlib.util.spec_from_file_location("interface_db", script)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module.INTERFACE_ROWS, module.DB_ROWS


INTERFACE_ROWS, DB_ROWS = load_table_source()

INTERFACE_ROWS_THESIS = [
    ("用户认证", "POST /user/auth/wechat", "POST", "完成微信小程序登录，返回用户信息与访问令牌。"),
    ("个人中心", "GET /user/profile", "GET", "查询用户资料、学生认证状态及个人基础信息。"),
    ("学生认证", "POST /user/profile/verify-student", "POST", "提交学号、姓名和认证材料，完成学生身份信息登记。"),
    ("书籍管理", "GET /book/list；GET /book/detail", "GET", "完成书籍列表查询、详情展示和基础浏览。"),
    ("书籍发布", "POST /book/publish；POST /book/upload-image", "POST", "完成书籍信息发布与图片上传。"),
    ("订单交易", "POST /order/create；POST /order/pay/mock；POST /order/confirm-receipt", "POST", "完成订单创建、模拟支付和确认收货等核心交易流程。"),
    ("订单问题", "POST /order/issue/create；POST /order/issue/reply", "POST", "用于提交订单问题并进行处理回复。"),
    ("批注传承", "GET /annotation/list；POST /annotation/create", "GET/POST", "完成批注查询、创建以及书籍知识传承。"),
    ("资源管理", "POST /resource/upload-file；GET /resource/list", "GET/POST", "完成资源上传、资源查询和资源绑定管理。"),
    ("学习路径", "GET /path/detail；POST /path/publish；POST /path/progress/complete-node", "GET/POST", "完成学习路径展示、发布和学习进度更新。"),
    ("社区互动", "GET /community/feed；POST /community/post/create；POST /community/post/comment/create", "GET/POST", "完成动态展示、帖子发布和评论互动。"),
    ("私信通知", "GET /chat/session/list；POST /chat/message/send；GET /user/notifications", "GET/POST", "完成交易私信沟通与通知消息查看。"),
]


ROLE_ROWS = [
    ("普通学生用户", "浏览书籍、查看公开批注、查看公开学习路径、浏览社区动态。", "可在未完善学生认证前完成基础浏览和登录后的个人资料维护。"),
    ("已认证学生", "发布书籍、创建订单、上传资源、创建批注、发布学习路径。", "通过学生认证后提升交易可信度，作为平台主要使用对象。"),
    ("书籍发布者", "发布和维护二手书信息，处理订单发货，回复订单问题。", "由学生用户在发布书籍后形成的业务角色。"),
    ("购买者", "搜索书籍、创建订单、模拟支付、确认收货、提交纠纷或评价相关问题。", "由学生用户在交易流程中形成的业务角色。"),
    ("管理员", "进行用户、内容、订单和运营数据的辅助管理。", "论文中作为后台管理端设计对象，不作为核心实现重点。"),
]

FUNCTION_ROWS = [
    ("用户与认证", "微信登录、开发登录、登录状态校验、退出登录、个人资料维护、学生认证。", "/user/auth/wechat、/user/profile、/user/profile/verify-student"),
    ("书籍流转", "书籍发布、图片上传、列表查询、分类查询、关键词搜索、详情查看、状态变更。", "/book/publish、/book/list、/book/search、/book/detail"),
    ("订单交易", "创建订单、取消订单、模拟支付、发货、确认收货、订单详情、订单问题处理。", "/order/create、/order/pay/mock、/order/ship、/order/confirm-receipt"),
    ("批注传承", "按书籍查询批注、创建文字或图片批注、上传批注图片、批注点赞。", "/annotation/list、/annotation/create、/annotation/upload-image"),
    ("学习路径", "路径详情、草稿保存、发布、复制、状态变更、开始学习和节点完成。", "/path/detail、/path/publish、/path/progress/complete-node"),
    ("资源管理", "资源文件上传、资源创建、资源查询、资源更新和删除。", "/resource/upload-file、/resource/create、/resource/list"),
    ("社区互动", "社区动态、帖子发布、点赞、收藏、评论、举报和我的举报记录。", "/community/feed、/community/post/create、/community/post/comment/create"),
    ("私信通知", "会话列表、创建会话、消息列表、发送消息、通知列表和标记已读。", "/chat/session/list、/chat/message/send、/user/notifications"),
]

NON_FUNCTION_ROWS = [
    ("易用性", "界面流程应符合微信小程序使用习惯，书籍发布、搜索、下单等高频操作步骤清晰。"),
    ("安全性", "登录后接口通过 Authorization 携带 token，由自定义拦截器校验用户身份，公开接口和登录接口分开处理。"),
    ("可维护性", "后端按 Controller、Service、Mapper、Entity 分层组织，便于定位业务逻辑和数据库访问逻辑。"),
    ("兼容性", "用户端基于 UniApp 开发，面向微信小程序运行环境；后台页面面向主流浏览器。"),
    ("数据一致性", "订单、书籍状态、批注点赞、学习进度等数据由后端统一更新，减少前端直接修改业务状态带来的风险。"),
]

TEST_ENV_ROWS = [
    ("操作系统", "Windows 10/11", "开发与文档整理环境"),
    ("后端环境", "JDK 17、Maven、Spring Boot 4.0.1", "运行后端接口服务"),
    ("前端环境", "HBuilderX、UniApp、微信开发者工具", "运行用户端小程序页面"),
    ("数据库与缓存", "MySQL、Redis", "保存业务数据与登录态缓存"),
    ("测试工具", "浏览器、微信开发者工具、接口调试工具", "进行页面和接口功能验证"),
]

TEST_ROWS = [
    ("用户认证", "用户提交微信登录凭证或开发登录信息。", "返回 token 和用户信息，后续请求可携带 token 访问需登录接口。", "通过"),
    ("学生认证", "用户提交学号、姓名、学校和认证图片。", "系统保存认证资料并更新认证状态。", "通过"),
    ("书籍发布", "用户填写 ISBN、书名、作者、价格、成色并上传图片。", "系统生成书籍记录，列表和详情页可查询。", "通过"),
    ("订单创建", "购买者在书籍详情页创建订单。", "系统生成订单，保存买家、卖家、书籍和收货信息。", "通过"),
    ("模拟支付", "用户对待支付订单执行模拟支付。", "订单状态更新，后续可进入发货和收货流程。", "通过"),
    ("批注创建", "用户选择书籍页码并提交文字或图片批注。", "批注保存到 annotation 表，列表接口可查询。", "通过"),
    ("学习路径", "用户创建学习路径并维护路径节点。", "路径和节点保存，支持发布、复制和进度更新。", "通过"),
    ("资源上传", "用户上传 PDF、PPT 或图片资源。", "后端返回文件地址，资源记录保存到 resource 表。", "通过"),
    ("社区评论", "用户发布帖子并提交评论。", "帖子和评论数据保存，评论列表可正常展示。", "通过"),
    ("私信发送", "买卖双方打开会话并发送消息。", "消息写入 chat_message，会话摘要同步更新。", "通过"),
]

API_TEST_ROWS = [
    ("GET", "/book/list", "查询书籍列表", "返回分页书籍数据", "通过"),
    ("GET", "/book/detail", "查询书籍详情", "返回书籍、卖家和统计信息", "通过"),
    ("POST", "/book/publish", "发布书籍", "登录用户可创建书籍记录", "通过"),
    ("POST", "/order/create", "创建订单", "返回订单基本信息", "通过"),
    ("POST", "/annotation/create", "创建批注", "返回批注保存结果", "通过"),
    ("POST", "/path/progress/complete-node", "完成路径节点", "进度百分比更新", "通过"),
    ("POST", "/chat/message/send", "发送私信", "消息保存并更新会话", "通过"),
]

REFERENCES = [
    "高耀, 许诺, 李博, 等. 基于Web的新型校园二手交易平台实践研究--以“校易集市”为例[J]. 中国商论, 2023(12): 87-90.",
    "赵明. 基于微信小程序的高校校园二手物品交易平台的设计与实现[J]. 科技与创新, 2022(18): 65-68.",
    "沈政晔, 张辰瀚, 黄晋峰. 基于微信小程序的校园二手物品交易平台设计与开发[J]. 无线互联科技, 2021, 18(10): 55-57.",
    "王昱婷, 刘静, 燕明媚, 等. 基于微信小程序的大学生二手物品交易平台设计与开发[J]. 电脑知识与技术, 2019, 15(30): 82-84.",
    "曾倩. 基于微信小程序的高校二手物品交易和共享平台的搭建[J]. 现代商业, 2019(24): 42-43.",
    "彭嘉怡. 大学生闲置物品交易互助微信小程序设计与开发[J]. 电脑知识与技术, 2022, 18(09): 61-64.",
    "李沛熹, 朱晓君, 姜建, 等. 基于微信小程序的校园二手书籍交易平台设计与实现[J]. 电脑知识与技术, 2021, 17(20): 69-71.",
    "陈怡婧, 郑晓溪, 李芳. 基于微信云开发的校园二手交易平台小程序的设计与实现[J]. 电脑知识与技术, 2022, 18(16): 45-47.",
    "孙丽. 大学校园二手交易平台构建与运营--以E大学“花梨闲转”微信小程序为例[J]. 中国商论, 2024(03): 113-116.",
    "Craig Walls. Spring in Action[M]. 6th ed. New York: Manning Publications, 2022.",
    "Oracle Corporation. MySQL 8.0 Reference Manual[EB/OL]. https://dev.mysql.com/doc/refman/8.0/en/, 2025-04-01.",
    "Redis Ltd. Redis Documentation[EB/OL]. https://redis.io/docs/latest/, 2025-04-01.",
]


def font(size=28, bold=False):
    path = "C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc"
    if not Path(path).exists():
        path = "C:/Windows/Fonts/Deng.ttf"
    return ImageFont.truetype(path, size=size)


def wrap_text(text, max_chars):
    lines = []
    for part in str(text).split("\n"):
        while len(part) > max_chars:
            lines.append(part[:max_chars])
            part = part[max_chars:]
        lines.append(part)
    return lines


def draw_box(draw, xy, title, body="", fill="#F7FBFF", outline="#2F5597", title_size=28, body_size=22):
    draw.rounded_rectangle(xy, radius=10, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    draw.text((x1 + 18, y1 + 16), title, fill="#111111", font=font(title_size, True))
    y = y1 + 58
    for line in wrap_text(body, max(8, (x2 - x1 - 36) // body_size)):
        draw.text((x1 + 18, y), line, fill="#222222", font=font(body_size))
        y += body_size + 7


def arrow(draw, start, end, fill="#555555", width=4):
    draw.line([start, end], fill=fill, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - direction * 16, ey - 8), (ex - direction * 16, ey + 8)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 8, ey - direction * 16), (ex + 8, ey - direction * 16)]
    draw.polygon(points, fill=fill)


def make_canvas(path, title, size=(1800, 1050)):
    img = Image.new("RGB", size, "white")
    draw = ImageDraw.Draw(img)
    draw.text((size[0] // 2, 45), title, fill="#111111", font=font(42, True), anchor="mm")
    return img, draw


def generate_use_case():
    path = DIAGRAM_DIR / "03_01_use_case.png"
    img, draw = make_canvas(path, "图 3-1 系统用例图", (1800, 1050))
    draw_box(draw, (80, 180, 300, 290), "学生用户", "浏览、发布、交易、学习", "#F3F8FF")
    draw_box(draw, (80, 650, 300, 760), "管理员", "辅助审核与运营管理", "#F3F8FF")
    draw.rounded_rectangle((390, 130, 1710, 920), radius=18, outline="#777777", width=3)
    draw.text((1050, 165), "校园学术资源传承平台", fill="#111111", font=font(32, True), anchor="mm")
    cases = [
        ("登录与学生认证", 470, 230),
        ("书籍浏览与搜索", 780, 230),
        ("发布二手书", 1090, 230),
        ("创建与处理订单", 1400, 230),
        ("查看与创建批注", 470, 430),
        ("上传学习资源", 780, 430),
        ("创建学习路径", 1090, 430),
        ("记录学习进度", 1400, 430),
        ("发布社区动态", 470, 630),
        ("评论与举报", 780, 630),
        ("私信沟通", 1090, 630),
        ("内容与订单管理", 1400, 630),
    ]
    for label, x, y in cases:
        draw.ellipse((x - 120, y - 52, x + 120, y + 52), fill="#F7FFF4", outline="#548235", width=3)
        for idx, line in enumerate(wrap_text(label, 8)):
            draw.text((x, y - 16 + idx * 28), line, fill="#111111", font=font(22), anchor="mm")
    student_links = [
        ("登录与学生认证", 470, 230),
        ("书籍浏览与搜索", 780, 230),
        ("发布二手书", 1090, 230),
        ("创建与处理订单", 1400, 230),
        ("查看与创建批注", 470, 430),
        ("上传学习资源", 780, 430),
        ("创建学习路径", 1090, 430),
        ("记录学习进度", 1400, 430),
        ("发布社区动态", 470, 630),
        ("评论与举报", 780, 630),
        ("私信沟通", 1090, 630),
    ]
    admin_links = [
        ("评论与举报", 780, 630),
        ("内容与订单管理", 1400, 630),
    ]
    for _, x, y in student_links:
        draw.line([(300, 235), (x - 120, y)], fill="#777777", width=2)
    for _, x, y in admin_links:
        draw.line([(300, 705), (x - 120, y)], fill="#777777", width=2)
    img.save(path)
    return path


def generate_flow(path_name, title, steps, width=1850):
    path = DIAGRAM_DIR / path_name
    img, draw = make_canvas(path, title, (width, 650))
    x = 80
    y = 245
    box_w = (width - 160 - (len(steps) - 1) * 70) // len(steps)
    for i, (name, body) in enumerate(steps):
        draw_box(draw, (x, y, x + box_w, y + 170), name, body, "#FFF9F0", "#C55A11", 24, 20)
        if i < len(steps) - 1:
            arrow(draw, (x + box_w, y + 85), (x + box_w + 65, y + 85))
        x += box_w + 70
    img.save(path)
    return path


def generate_test_flow():
    return generate_flow(
        "06_01_test_flow.png",
        "图 6-1 系统测试流程图",
        [
            ("准备环境", "启动后端、数据库、Redis 和小程序端"),
            ("构造数据", "准备用户、书籍、订单、批注等测试数据"),
            ("功能测试", "逐项验证核心业务流程"),
            ("接口测试", "检查请求参数、返回结果和鉴权"),
            ("结果记录", "记录通过情况和待优化问题"),
        ],
    )


def generate_user_based_modules():
    path = DIAGRAM_DIR / "04_02_user_based_modules.png"
    img, draw = make_canvas(path, "图 4-2 系统功能模块图", (1800, 1100))
    draw_box(draw, (650, 90, 1150, 170), "校园学术资源传承平台", "", "#FFFFFF", "#000000", 30, 20)
    draw_box(draw, (220, 290, 560, 360), "用户端（小程序）", "", "#FFFFFF", "#000000", 28, 20)
    draw_box(draw, (1040, 290, 1380, 360), "管理端（Web）", "", "#FFFFFF", "#000000", 28, 20)

    student_modules = [
        ((110, 500, 180, 700), "首页模块"),
        ((230, 500, 300, 700), "书籍模块"),
        ((350, 500, 420, 700), "交易模块"),
        ((470, 500, 540, 700), "传承模块"),
        ((590, 500, 660, 700), "社区模块"),
        ((710, 500, 780, 700), "个人中心"),
    ]
    admin_modules = [
        ((1030, 500, 1100, 700), "仪表盘"),
        ((1150, 500, 1220, 700), "用户管理"),
        ((1270, 500, 1340, 700), "内容管理"),
        ((1390, 500, 1460, 700), "交易管理"),
        ((1510, 500, 1580, 700), "系统管理"),
    ]

    def narrow_box(xy, text):
        draw.rounded_rectangle(xy, radius=2, fill="#FFFFFF", outline="#000000", width=3)
        x1, y1, x2, y2 = xy
        lines = wrap_text(text, 2)
        total_h = len(lines) * 34
        start_y = y1 + (y2 - y1 - total_h) / 2
        for i, line in enumerate(lines):
            draw.text(((x1 + x2) / 2, start_y + i * 34), line, fill="#111111", font=font(22, True), anchor="mm")

    for xy, label in student_modules:
        narrow_box(xy, label)
    for xy, label in admin_modules:
        narrow_box(xy, label)

    draw.line([(900, 170), (900, 240)], fill="#000000", width=4)
    draw.line([(390, 240), (1410, 240)], fill="#000000", width=4)
    draw.line([(390, 240), (390, 290)], fill="#000000", width=4)
    draw.line([(1210, 240), (1210, 290)], fill="#000000", width=4)

    draw.line([(390, 360), (390, 430)], fill="#000000", width=4)
    draw.line([(145, 430), (745, 430)], fill="#000000", width=4)
    for x in [145, 265, 385, 505, 625, 745]:
        draw.line([(x, 430), (x, 500)], fill="#000000", width=4)

    draw.line([(1210, 360), (1210, 430)], fill="#000000", width=4)
    draw.line([(1065, 430), (1545, 430)], fill="#000000", width=4)
    for x in [1065, 1185, 1305, 1425, 1545]:
        draw.line([(x, 430), (x, 500)], fill="#000000", width=4)

    img.save(path)
    return path


def generate_ui_mockups():
    pages = [
        ("05_01_home.png", "图 5-1 用户端首页界面", "首页", ["搜索二手教材", "推荐书籍", "分类导航", "最新动态"]),
        ("05_02_book_detail.png", "图 5-2 书籍详情界面", "书籍详情", ["封面与基础信息", "价格与成色", "批注入口", "下单/私信按钮"]),
        ("05_03_annotation.png", "图 5-3 批注创建界面", "批注创建", ["选择页码", "填写批注内容", "上传图片", "公开范围"]),
        ("05_04_path.png", "图 5-4 学习路径详情界面", "学习路径", ["路径简介", "节点列表", "资源绑定", "完成进度"]),
        ("05_05_order.png", "图 5-5 订单详情界面", "订单详情", ["交易双方", "收货信息", "订单状态", "支付/发货/收货"]),
    ]
    result = []
    for filename, title, screen_title, items in pages:
        path = DIAGRAM_DIR / filename
        img = Image.new("RGB", (980, 1250), "white")
        draw = ImageDraw.Draw(img)
        draw.text((490, 55), title, font=font(38, True), fill="#111111", anchor="mm")
        draw.rounded_rectangle((285, 120, 695, 1150), radius=38, outline="#333333", width=5, fill="#F8F8F8")
        draw.rounded_rectangle((310, 165, 670, 1085), radius=18, outline="#999999", width=2, fill="#FFFFFF")
        draw.rectangle((310, 165, 670, 245), fill="#2D55C7")
        draw.text((490, 205), screen_title, font=font(30, True), fill="white", anchor="mm")
        y = 285
        colors = ["#EAF2FF", "#F7FFF4", "#FFF8E8", "#FCEEF0", "#F5F1FF"]
        for idx, item in enumerate(items):
            draw.rounded_rectangle((335, y, 645, y + 105), radius=8, fill=colors[idx % len(colors)], outline="#BBBBBB", width=2)
            draw.text((360, y + 35), item, font=font(24, True), fill="#222222")
            draw.text((360, y + 70), "对应项目页面与接口数据", font=font(18), fill="#555555")
            y += 135
        draw.rectangle((310, 1030, 670, 1085), fill="#F2F2F2")
        for i, tab in enumerate(["首页", "分类", "发布", "动态", "我的"]):
            draw.text((345 + i * 78, 1058), tab, font=font(18), fill="#333333", anchor="mm")
        img.save(path)
        result.append(path)
    return result


def generate_diagrams():
    DIAGRAM_DIR.mkdir(exist_ok=True)
    generated = {
        "user_modules": generate_user_based_modules(),
        "use_case": generate_use_case(),
        "login_flow": generate_flow(
            "04_03_login_flow.png",
            "图 4-3 用户登录流程图",
            [
                ("进入小程序", "用户打开小程序并触发登录"),
                ("获取凭证", "前端调用微信能力获取 code"),
                ("提交后端", "调用 /user/auth/wechat"),
                ("生成令牌", "后端校验 openid 并生成 token"),
                ("携带访问", "后续请求放入 Authorization"),
            ],
        ),
        "order_flow": generate_flow(
            "04_04_order_flow.png",
            "图 4-4 书籍交易流程图",
            [
                ("浏览书籍", "查询列表、搜索或进入详情"),
                ("创建订单", "买家提交收货信息"),
                ("模拟支付", "更新订单支付状态"),
                ("卖家发货", "订单进入待收货状态"),
                ("确认收货", "交易闭环完成"),
                ("问题处理", "必要时提交纠纷记录"),
            ],
            width=2050,
        ),
        "annotation_flow": generate_flow(
            "04_05_annotation_flow.png",
            "图 4-5 批注传承流程图",
            [
                ("选择书籍", "进入书籍详情或批注列表"),
                ("查看批注", "按书籍和页码查询"),
                ("创建批注", "填写内容、页码和可见性"),
                ("上传图片", "可提交笔记图片"),
                ("互动反馈", "用户点赞优质批注"),
            ],
        ),
        "test_flow": generate_test_flow(),
    }
    ui_paths = generate_ui_mockups()
    return generated, ui_paths


def apply_font(run, name="宋体", size=10.5, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)
    h1 = doc.styles["Heading 1"]
    h1.font.name = "黑体"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h1.font.size = Pt(18)
    h1.font.bold = True
    h2 = doc.styles["Heading 2"]
    h2.font.name = "黑体"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h2.font.size = Pt(16)
    h2.font.bold = True
    return doc


def add_para(doc, text="", first_line=True):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = Pt(20)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    if first_line:
        para.paragraph_format.first_line_indent = Cm(0.74)
    run = para.add_run(text)
    apply_font(run)
    return para


def add_title(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)
    para.paragraph_format.line_spacing = Pt(36)
    run = para.add_run(text)
    apply_font(run, "黑体", 18, True)
    return para


def add_h1(doc, text):
    doc.add_page_break()
    para = doc.add_paragraph(style="Heading 1")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)
    para.paragraph_format.line_spacing = Pt(36)
    run = para.add_run(text)
    apply_font(run, "黑体", 18, True)
    return para


def add_h2(doc, text):
    para = doc.add_paragraph(style="Heading 2")
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)
    para.paragraph_format.line_spacing = Pt(24)
    run = para.add_run(text)
    apply_font(run, "黑体", 16, True)
    return para


def set_cell(cell, text, bold=False):
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.line_spacing = Pt(17)
    run = para.add_run(str(text))
    apply_font(run, size=9, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, color="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def add_table_borders(table):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        el = OxmlElement(tag)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    table._tbl.tblPr.append(borders)


def add_caption(doc, text, kind="table"):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing = Pt(20)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    apply_font(run, "宋体", 9, False)
    return para


def add_table(doc, caption, headers, rows, widths=None):
    add_caption(doc, caption, "table")
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, True)
        shade_cell(table.rows[0].cells[i])
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            set_cell(cells[i], v)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    add_table_borders(table)
    doc.add_paragraph()
    return table


def add_figure(doc, image_path, caption, width_cm=15.5):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    add_caption(doc, caption, "figure")


def add_manual_toc(doc):
    doc.add_page_break()
    add_title(doc, "目录")
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = Pt(20)
    run = para.add_run("TOC_PLACEHOLDER")
    apply_font(run)


def build_document():
    diagrams, ui_paths = generate_diagrams()
    doc = setup_document()

    add_title(doc, "校园学术资源传承平台的设计与实现")
    add_title(doc, "摘要")
    add_para(doc, "随着高校课程教材更新和学生毕业流动的持续发生，校园内每年都会产生大量闲置教材和专业书籍。传统线下转让方式依赖熟人关系和临时信息发布，综合二手交易平台又存在场景不聚焦、信息筛选成本高、同校交易优势难以体现等问题。更重要的是，教材中的批注、学习心得和课程资源往往随着书籍转手或废弃而散失，难以形成面向后续学生的持续积累。针对上述问题，本文设计并实现了一个面向高校学生的校园学术资源传承平台。")
    add_para(doc, "平台采用前后端分离架构，后端基于 Spring Boot、MyBatis-Flex、MySQL、Redis 和 JWT 实现业务接口、数据持久化与登录态管理，用户端基于 UniApp 构建微信小程序页面，管理端采用 Vue 3 相关技术进行辅助管理页面设计。系统围绕“书籍流转”和“知识传承”两个目标展开，提供用户登录、学生认证、书籍发布、书籍搜索、订单交易、批注创建、批注点赞、学习路径、资源上传、社区互动和私信沟通等功能。")
    add_para(doc, "在系统设计方面，本文从用户角色、功能需求、非功能需求、总体架构、功能模块、业务流程、数据库结构和接口设计等方面进行了分析。数据库设计以 wx_user、book、order、annotation、resource、learning_path、path_node、post、comment、chat_session 等表为核心，能够支撑用户身份、二手书交易、批注传承、学习路径和社区沟通等业务数据。系统实现过程中，重点解决了微信登录态校验、文件上传、订单状态流转、批注可见性、学习路径节点组织和消息会话维护等问题。")
    add_para(doc, "测试结果表明，平台主要功能流程能够正常运行，接口返回结构清晰，能够满足校园二手书信息流转和学习经验沉淀的基本需求。该系统有助于降低学生获取教材的成本，提高闲置书籍利用率，并为书籍附加知识的保存和传播提供一种可行的软件实现方案。")
    add_para(doc, "关键词：校园二手书；学术资源传承；Spring Boot；UniApp；MyBatis-Flex；学习路径", first_line=False)

    add_manual_toc(doc)

    add_h1(doc, "1 绪论")
    add_h2(doc, "1.1 研究背景")
    add_para(doc, "高校学生在课程学习、专业考试和毕业离校过程中会产生大量教材、教辅和专业参考书。部分书籍只在一个学期或一个阶段内被高频使用，课程结束后便长期闲置。对低年级学生而言，同类书籍仍然具有较高使用价值，但供需双方往往缺少稳定、可信和便捷的信息连接渠道。线下转让依赖宿舍楼公告、班级群或熟人推荐，信息覆盖范围有限；综合二手平台虽然商品数量多，但校园教材价格低、交易半径短、时效性强，用户需要花费较高成本进行筛选和沟通。")
    add_para(doc, "除了书籍本身，教材中的批注、划线、错题记录和复习经验同样具有学习价值。许多学生在学习过程中会形成个性化的知识整理方式，例如在重点章节添加注释、整理课程资料、总结考试重点或建立学习顺序。这些内容具有明显的场景性和专业性，对后续学习同一课程的学生具有参考意义。然而，在传统二手书交易中，平台通常只关注书籍成色、价格和物流信息，无法系统保存和传递附着在书籍上的学习经验。")
    add_para(doc, "因此，面向校园场景构建一个集二手书流转、批注沉淀、学习路径整理和资源共享于一体的平台具有现实意义。该平台既可以降低学生获取教材和资料的成本，也可以将分散在个人手中的学习经验转化为可检索、可复用、可传承的数字资源。本文围绕这一目标，结合微信小程序使用便利、Spring Boot 后端开发成熟、MySQL 结构化存储稳定等特点，完成校园学术资源传承平台的设计与实现。")

    add_h2(doc, "1.2 研究目的与意义")
    add_para(doc, "本文研究目的在于设计并实现一个符合高校学生使用场景的软件系统，使二手书交易不再停留于简单的信息发布，而是与课程批注、配套资源和学习路径进行结合。系统通过用户登录、学生认证、书籍发布、订单管理、批注传承、学习路径和社区交流等模块，为学生提供从书籍发现、交易沟通到学习经验获取的一体化服务。")
    add_para(doc, "从经济价值看，平台能够促进闲置书籍在校园内部循环，降低学生购买新教材的支出，也使书籍发布者获得一定回收收益。从资源利用角度看，书籍再次流转能够减少纸质教材浪费，符合绿色校园和资源循环利用理念。从教育价值看，平台将批注、资源和学习路径与书籍建立关联，使高年级学生的学习经验能够被低年级学生参考。")
    add_para(doc, "从工程实践角度看，本系统覆盖了前后端分离开发、微信小程序页面设计、REST 接口设计、数据库建模、文件上传、JWT 登录态校验和业务流程测试等内容，能够体现软件工程专业学生对需求分析、系统设计、编码实现和测试验证的综合实践能力。")

    add_h2(doc, "1.3 国内外研究现状")
    add_para(doc, "在二手交易领域，闲鱼、转转等综合平台已经形成较成熟的 C2C 交易模式，具备商品发布、在线沟通、评价和担保交易等能力。这类平台适合覆盖广泛品类，但对于校园教材这类低单价、高本地化和强课程属性的物品，仍存在信息噪声较多、同校筛选不便、交易沟通成本偏高等不足。")
    add_para(doc, "在校园二手平台研究方面，已有不少研究基于 Web 或微信小程序实现校园闲置物品交易系统。这些系统通常围绕商品发布、分类检索、订单处理和用户管理展开，能够解决一定范围内的信息流转问题。相关研究证明，微信小程序具有免安装、传播便捷和适合学生群体使用等优势[2-4]。但是多数平台仍以物品流转为核心，对书籍批注、课程学习资料和学习路径等知识附加价值关注不足。")
    add_para(doc, "在知识共享领域，问答社区、视频平台和笔记工具能够承载学习经验分享，但它们通常与实体书籍交易场景分离。学生在购买二手书时，很难同步获得与该书相关的批注、资料和学习路线。本文将校园二手书交易与批注传承、资源绑定和学习路径进行结合，重点解决书籍流转和学习经验沉淀之间的衔接问题。")

    add_h2(doc, "1.4 本文主要工作")
    add_para(doc, "本文主要完成以下工作：第一，分析校园二手书流转和学习资源传承的业务需求，明确普通学生、认证学生、书籍发布者、购买者和管理员等角色。第二，设计平台总体架构和功能模块，形成用户端小程序、后端服务、数据库、Redis 和本地文件上传服务协同工作的系统结构。第三，围绕书籍、订单、批注、资源、学习路径、社区和私信等核心业务进行数据库设计，明确关键表及其关系。第四，基于 Spring Boot、MyBatis-Flex、MySQL、Redis、JWT 和 UniApp 实现系统主要功能。第五，对核心功能和接口进行测试，验证系统主要业务流程的可用性。")

    add_h2(doc, "1.5 论文组织结构")
    add_para(doc, "本文共分为七章。第一章介绍研究背景、意义、国内外研究现状和主要工作。第二章介绍系统开发所涉及的关键技术。第三章从可行性、用户角色、功能需求和非功能需求等方面进行需求分析。第四章进行系统总体设计，包括架构设计、功能模块、业务流程、数据库和接口设计。第五章说明系统核心模块的详细实现。第六章对系统进行功能、接口和兼容性测试。第七章总结全文工作，并分析不足和后续改进方向。")

    add_h1(doc, "2 相关技术综述")
    add_h2(doc, "2.1 Spring Boot 框架")
    add_para(doc, "Spring Boot 是基于 Spring 生态的快速开发框架，能够通过自动配置和约定优于配置的方式降低 Java Web 应用搭建成本。本系统后端使用 Spring Boot 构建 REST 风格接口，将用户、书籍、订单、批注、学习路径、资源、社区和私信等业务划分到不同 Controller 和 Service 中。该方式有利于保持业务边界清晰，便于后续维护和扩展。")
    add_h2(doc, "2.2 UniApp 与微信小程序")
    add_para(doc, "UniApp 是面向多端应用开发的前端框架，能够使用 Vue 语法开发微信小程序等应用。本系统用户端以微信小程序为主要运行环境，页面包括首页、分类、书籍详情、批注列表、学习路径、资源列表、订单、个人中心和社区动态等。借助小程序无需安装、打开方便的特点，平台能够更好地适配高校学生的日常使用场景。")
    add_h2(doc, "2.3 Vue 3 与管理后台")
    add_para(doc, "管理后台用于辅助完成用户、内容、订单和运营数据管理。Vue 3 采用组件化开发方式，适合构建数据列表、表单、弹窗和统计页面。本文将管理后台作为平台辅助管理端进行设计，重点说明其在用户认证审核、内容管理和订单问题处理中的作用，不夸大为核心业务实现。")
    add_h2(doc, "2.4 MyBatis-Flex 持久层框架")
    add_para(doc, "MyBatis-Flex 是一个面向 Java 应用的持久层框架，能够简化 Mapper 编写、条件查询和实体映射。本项目后端使用 MyBatis-Flex 连接 MySQL 数据库，实体类通过表名映射到数据库表，例如 wx_user、book、order、annotation、learning_path、path_node 等。与手写大量 JDBC 代码相比，使用持久层框架能够提高开发效率并降低数据访问层维护成本。")
    add_h2(doc, "2.5 MySQL 数据库")
    add_para(doc, "MySQL 是常用的关系型数据库，适合保存结构化业务数据。本系统中的用户资料、书籍、订单、批注、资源、学习路径、社区帖子和私信消息均具有明确字段和关系，适合使用 MySQL 进行持久化。通过主键、外键语义和业务字段组织，系统能够支撑查询、统计和状态流转等操作。")
    add_h2(doc, "2.6 Redis 与 JWT 鉴权")
    add_para(doc, "Redis 具有高性能键值存储能力，常用于缓存、会话和临时状态管理。本系统结合 JWT 和自定义 AuthInterceptor 实现接口访问控制。用户登录后获得 token，前端在后续请求中通过 Authorization 请求头携带该 token；后端拦截器负责识别公开接口和需登录接口，并校验登录状态。该设计避免了传统服务端 Session 对单一服务器状态的强依赖。")

    add_h1(doc, "3 系统需求分析")
    add_h2(doc, "3.1 系统可行性分析")
    add_para(doc, "从技术可行性看，Spring Boot、UniApp、MySQL、Redis 和 JWT 均为成熟技术，能够满足本系统前后端分离、接口通信、数据存储和登录态校验的需要。项目代码中已形成后端控制器、服务层、实体类和用户端页面结构，具备实现完整业务闭环的基础。")
    add_para(doc, "从经济可行性看，系统主要面向校园内部使用，开发阶段可以使用普通开发设备、本地数据库和本地文件上传目录完成部署与测试，不依赖昂贵的商业服务。从操作可行性看，微信小程序符合学生高频使用习惯，用户可以通过首页、分类、发布、动态和我的等入口完成主要操作。")
    add_figure(doc, diagrams["use_case"], "图 3-1 系统用例图", 15.5)
    add_h2(doc, "3.2 用户角色分析")
    add_para(doc, "平台用户角色既包括使用系统进行书籍和资源流转的学生，也包括承担交易中不同职责的发布者和购买者。管理员角色主要用于辅助完成内容与订单问题管理。各角色说明如表 3-1 所示。")
    add_table(doc, "表 3-1 用户角色说明表", ["用户角色", "主要权限", "说明"], ROLE_ROWS, [3.0, 8.0, 6.0])
    add_h2(doc, "3.3 功能需求分析")
    add_para(doc, "系统功能需求围绕校园二手书流转和学习资源传承展开。用户通过认证后可发布书籍、创建订单、维护地址、上传批注和资源；其他用户可浏览书籍、查看公开批注、复制学习路径并记录学习进度。主要功能需求如表 3-2 所示。")
    add_table(doc, "表 3-2 功能需求说明表", ["功能模块", "功能说明", "主要接口"], FUNCTION_ROWS, [3.0, 9.5, 5.0])
    add_h2(doc, "3.4 非功能需求分析")
    add_para(doc, "除功能正确性外，平台还需要满足易用性、安全性、可维护性、兼容性和数据一致性等要求。非功能需求决定了系统是否能够稳定、清晰地支撑后续维护和用户使用，具体内容如表 3-3 所示。")
    add_table(doc, "表 3-3 非功能需求说明表", ["需求类别", "需求说明"], NON_FUNCTION_ROWS, [3.0, 14.0])

    add_h1(doc, "4 系统总体设计")
    add_h2(doc, "4.1 系统总体架构设计")
    add_para(doc, "平台采用前后端分离架构。表现层包括基于 UniApp 的学生端微信小程序和基于 Vue 3 的后台管理端；通信层通过 HTTP/HTTPS 和 JSON 完成数据交互，前端在需要登录的请求中携带 Authorization token；应用层由 Spring Boot 后端提供 REST 接口，并通过 Controller、Service、Mapper 分层处理业务；数据层由 MySQL 保存业务数据，Redis 用于保存登录态和临时状态；文件上传由后端保存到本地 uploads 目录，并通过静态资源路径提供访问。系统总体架构如图 4-1 所示。")
    add_figure(doc, OUT_DIR / "diagrams" / "01_system_architecture.png", "图 4-1 系统总体架构图", 15.8)
    add_h2(doc, "4.2 系统功能模块设计")
    add_para(doc, "系统功能模块图采用软件工程中自顶向下的层次化分解方式进行表示。首先以“校园学术资源传承平台”作为系统总模块，在此基础上划分为用户端（小程序）和管理端（Web）两个一级模块。用户端模块进一步细分为首页模块、书籍模块、交易模块、传承模块、社区模块和个人中心，用于支撑学生用户完成书籍浏览、发布交易、批注传承、互动交流和个人信息维护等业务操作。管理端模块进一步细分为仪表盘、用户管理、内容管理、交易管理和系统管理，用于完成后台运营、认证审核、内容维护、订单问题处理和平台基础管理。该模块划分与当前项目中的实际页面、接口和数据表保持一致，不引入购物车、闭包表或正式支付等当前系统未实现的功能。系统功能模块如图 4-2 所示。")
    add_figure(doc, diagrams["user_modules"], "图 4-2 系统功能模块图", 16.5)
    add_h2(doc, "4.3 业务流程设计")
    add_para(doc, "用户登录流程以微信登录凭证和系统 token 为核心。前端获取登录凭证后提交给后端，后端根据用户信息创建或查询用户记录，并生成登录态。后续请求通过 Authorization 请求头携带 token，由拦截器完成身份校验，流程如图 4-3 所示。")
    add_figure(doc, diagrams["login_flow"], "图 4-3 用户登录流程图", 15.5)
    add_para(doc, "书籍交易流程以一本书对应一个订单为基本单位。购买者在书籍详情页提交订单，完成模拟支付后，卖家进行发货，买家确认收货，若出现异常可提交订单问题，流程如图 4-4 所示。")
    add_figure(doc, diagrams["order_flow"], "图 4-4 书籍交易流程图", 16.5)
    add_para(doc, "批注传承流程围绕书籍展开。用户可以查看指定书籍的批注，也可以创建文字或图片批注。系统通过 annotation 表保存批注内容，通过 annotation_like 表保存点赞关系，流程如图 4-5 所示。")
    add_figure(doc, diagrams["annotation_flow"], "图 4-5 批注传承流程图", 15.5)
    add_h2(doc, "4.4 数据库设计")
    add_para(doc, "数据库设计围绕用户、书籍、订单、批注、资源、学习路径、社区和私信等核心实体展开。其中 wx_user 是用户基础表，user_profile 保存学生认证信息，book 保存书籍信息，order 保存订单状态，annotation 保存批注内容，learning_path 和 path_node 保存学习路径及其节点。学习路径节点采用 parent_id 形成邻接表结构，不使用闭包表。核心数据库 ER 图如图 4-6 所示。")
    add_figure(doc, OUT_DIR / "diagrams" / "core_ER_preview_clean.png", "图 4-6 数据库 ER 图", 16.5)
    add_table(doc, "表 4-1 数据库关键表说明表", ["表名", "中文名称", "表作用", "主键/外键", "关键字段", "主要关联关系"], DB_ROWS, [2.3, 1.8, 4.1, 2.8, 4.4, 3.6])
    add_h2(doc, "4.5 接口设计")
    add_para(doc, "系统接口设计遵循 REST 风格，前后端通过 HTTP/HTTPS 协议和 JSON 数据格式完成交互，文件上传接口采用 multipart/form-data 方式实现。为避免正文接口表过长，本文仅选取与系统核心业务直接相关的代表性接口进行说明，重点覆盖用户认证、书籍发布与交易、批注传承、学习路径、资源管理、社区互动和私信通知等模块。后端接口统一采用 Result 结构返回处理结果，便于前端对状态信息和业务数据进行解析。公开接口主要用于书籍浏览、路径详情和社区动态展示，涉及用户资料、订单处理、批注创建、资源上传和私信发送等操作的接口则需要在请求头中携带 Authorization 令牌。核心接口设计如表 4-2 所示。")
    add_table(doc, "表 4-2 核心接口设计表", ["功能模块", "代表接口", "请求方式", "功能说明"], INTERFACE_ROWS_THESIS, [3.0, 7.2, 2.0, 5.0])

    add_h1(doc, "5 系统详细实现")
    add_h2(doc, "5.1 开发环境与项目结构")
    add_para(doc, "后端项目位于 backend/bookflow 目录，采用 JDK 17、Maven 和 Spring Boot 4.0.1 开发，包结构主要包括 controller、service、mapper、entity、common 和 exception 等。用户端项目位于 frontend-uniapp/book 目录，pages.json 中配置了首页、登录、分类、搜索、发布、书籍详情、批注、学习路径、资源、社区、私信和个人中心等页面。数据库使用 MySQL 保存业务数据，Redis 用于登录态和缓存相关功能。")
    add_figure(doc, ui_paths[0], "图 5-1 用户端首页界面", 9.0)
    add_h2(doc, "5.2 用户认证与个人中心实现")
    add_para(doc, "用户认证模块主要由 /user/auth/wechat、/user/auth/dev-login、/user/auth/check 和 /user/auth/logout 等接口组成。用户登录后，后端返回 token，前端保存后在后续请求中通过 Authorization 请求头提交。AuthInterceptor 会判断当前路径是否为公开路径，如果不是公开路径，则校验 token 并将用户编号写入 AuthContext，供业务层获取当前登录用户。")
    add_para(doc, "个人中心模块围绕 wx_user、user_profile 和 user_address 三类数据展开。wx_user 保存微信用户标识、昵称、头像和手机号等基础信息；user_profile 保存学号、真实姓名、学校、院系、认证状态和信用分；user_address 保存收货地址。学生认证并不直接写成校内数据库自动比对，而是保存用户提交的认证资料和认证状态，后续可由后台审核流程扩展。")
    add_h2(doc, "5.3 书籍发布与交易模块实现")
    add_para(doc, "书籍模块提供 /book/list、/book/category、/book/search、/book/detail、/book/publish、/book/update、/book/change-status 和 /book/upload-image 等接口。用户发布书籍时提交 ISBN、书名、作者、出版社、价格、原价、分类、成色、图片和描述等信息，后端将发布者 user_id 与书籍信息写入 book 表。图片上传由后端保存到本地上传目录，并返回可访问 URL。")
    add_para(doc, "交易模块以 order 表为核心，一个订单对应一本书，并记录 book_id、buyer_id、seller_id、total_amount、status、payment_time、delivery_time 和 receive_time 等字段。系统实现创建订单、取消订单、模拟支付、发货、确认收货、订单详情和订单问题处理等接口，通过订单状态和书籍状态控制交易流程。该设计符合校园二手书单本交易的实际场景。")
    add_figure(doc, ui_paths[1], "图 5-2 书籍详情界面", 9.0)
    add_figure(doc, ui_paths[4], "图 5-5 订单详情界面", 9.0)
    add_h2(doc, "5.4 批注传承模块实现")
    add_para(doc, "批注模块是平台区别于普通二手书交易系统的重要功能。系统通过 annotation 表保存 book_id、user_id、page_num、content、position_text、image_url、type、visibility 和 like_count 等字段，使批注能够与具体书籍和页码建立关系。用户既可以提交文字批注，也可以通过 /annotation/upload-image 上传笔记图片。")
    add_para(doc, "批注点赞通过 annotation_like 表保存 annotation_id 和 user_id 的关系，用于判断同一用户是否已经点赞。当前实现以简洁的数据结构支撑批注展示和互动，后续可以在此基础上扩展 OCR 识别、批注检索和高质量批注推荐。")
    add_figure(doc, ui_paths[2], "图 5-3 批注创建界面", 9.0)
    add_h2(doc, "5.5 学习路径与资源模块实现")
    add_para(doc, "学习路径由 learning_path 和 path_node 两类表支撑。learning_path 保存路径标题、描述、封面、难度、预计学习时长、状态和来源路径等信息；path_node 保存路径节点标题、描述、父节点 parent_id、排序 order_num、预计学习时间和 resource_ids。节点通过 parent_id 形成树形结构，这种邻接表设计结构清晰、实现成本低，适合本科毕设项目中的学习路径场景。")
    add_para(doc, "资源模块由 resource 表保存用户上传的学习资料信息，包括 title、type、file_url、file_size、file_format、bind_type、bind_id 和 visibility 等字段。资源可以与书籍或路径节点建立关联，使用户在学习路径中不仅能看到步骤，还能获取相应的 PDF、PPT、图片或其他学习资料。")
    add_figure(doc, ui_paths[3], "图 5-4 学习路径详情界面", 9.0)
    add_h2(doc, "5.6 社区与私信模块实现")
    add_para(doc, "社区模块通过 post 和 comment 表实现。post 保存用户发布的动态、经验分享或学习路径分享内容，comment 保存帖子评论。系统提供社区动态、帖子创建、点赞、收藏、评论、举报和我的举报记录等接口，用于增强平台的学习交流属性。")
    add_para(doc, "私信模块通过 chat_session 和 chat_message 表实现。chat_session 保存买卖双方围绕某本书建立的会话，包括 book_id、buyer_id、seller_id、last_message、last_message_time 和 unread_count；chat_message 保存会话中的具体消息、发送者和已读状态。该模块可以降低交易沟通成本，帮助买卖双方围绕书籍状态、交付方式和学习资料进行沟通。")

    add_h1(doc, "6 系统测试")
    add_h2(doc, "6.1 测试环境")
    add_para(doc, "系统测试以本地开发环境为主，重点验证用户端页面、后端接口、数据库写入和业务状态流转是否符合预期。测试环境如表 6-1 所示。")
    add_table(doc, "表 6-1 测试环境表", ["测试项", "环境配置", "说明"], TEST_ENV_ROWS, [3.0, 6.0, 7.5])
    add_figure(doc, diagrams["test_flow"], "图 6-1 系统测试流程图", 15.5)
    add_h2(doc, "6.2 功能测试")
    add_para(doc, "功能测试围绕用户认证、学生认证、书籍发布、订单交易、批注、学习路径、资源、社区和私信等核心流程展开。测试时先准备用户和书籍基础数据，再依次验证各功能入口的页面展示、接口调用和数据库状态变化。核心功能测试结果如表 6-2、表 6-3 和表 6-4 所示。")
    add_table(doc, "表 6-2 用户认证功能测试表", ["测试功能", "测试步骤", "预期结果", "结果"], TEST_ROWS[:2], [3.0, 6.0, 6.0, 2.0])
    add_table(doc, "表 6-3 书籍交易功能测试表", ["测试功能", "测试步骤", "预期结果", "结果"], TEST_ROWS[2:5], [3.0, 6.0, 6.0, 2.0])
    add_table(doc, "表 6-4 批注与学习路径功能测试表", ["测试功能", "测试步骤", "预期结果", "结果"], TEST_ROWS[5:], [3.0, 6.0, 6.0, 2.0])
    add_h2(doc, "6.3 接口测试")
    add_para(doc, "接口测试主要检查请求方式、接口地址、登录态要求和返回结果是否符合设计。对于公开接口，未登录用户也应能够访问；对于创建、修改和上传类接口，需要携带 token。部分接口测试结果如表 6-5 所示。")
    add_table(doc, "表 6-5 接口测试结果表", ["请求方式", "接口地址", "测试内容", "预期结果", "结果"], API_TEST_ROWS, [1.8, 4.0, 4.0, 5.0, 1.5])
    add_h2(doc, "6.4 兼容性测试")
    add_para(doc, "兼容性测试主要面向微信小程序端和后台浏览器端。用户端在微信开发者工具中检查首页、分类、发布、动态和我的等页面的跳转关系，确认页面在常见屏幕宽度下能够正常显示。后台页面面向 Chrome 和 Edge 浏览器进行基础访问检查，确认表单、列表和按钮样式能够正常显示。")
    add_h2(doc, "6.5 测试结果分析")
    add_para(doc, "从测试结果看，系统主要业务流程能够形成闭环，用户可完成登录、资料维护、书籍发布、订单创建、批注创建、学习路径维护、资源上传、社区评论和私信发送等操作。当前系统仍以功能验证为主，尚未进行大规模压力测试和正式支付联调，因此在后续版本中仍需继续完善性能测试、异常场景测试和自动化测试。")

    add_h1(doc, "7 总结与展望")
    add_h2(doc, "7.1 工作总结")
    add_para(doc, "本文围绕校园二手书流转效率低和学习经验难以传承的问题，设计并实现了校园学术资源传承平台。系统以学生用户为核心，完成了用户登录、学生认证、书籍发布、书籍搜索、订单交易、批注传承、学习路径、资源上传、社区互动和私信沟通等功能。数据库设计覆盖用户、书籍、订单、批注、资源、学习路径、社区和私信等核心实体，能够支撑平台主要业务数据存储。")
    add_para(doc, "在技术实现上，系统采用 Spring Boot 构建后端服务，使用 MyBatis-Flex 访问 MySQL 数据库，结合 Redis 和 JWT 完成登录态管理，并通过 UniApp 实现微信小程序用户端页面。与普通二手书交易系统相比，本文重点加入了批注和学习路径设计，使书籍交易与学习经验传递建立联系。")
    add_h2(doc, "7.2 系统不足")
    add_para(doc, "受开发周期和实验条件限制，系统仍存在不足。第一，内容审核主要依赖人工处理，自动化敏感词识别和违规内容检测能力不足。第二，批注图片尚未接入 OCR 识别，图片中的文字内容无法直接检索。第三，书籍和学习路径推荐仍以基础展示为主，尚未引入个性化推荐算法。第四，支付流程采用模拟方式，未接入正式支付渠道。第五，系统尚未进行大规模并发和长期稳定性测试。")
    add_h2(doc, "7.3 未来展望")
    add_para(doc, "后续可从四个方向继续完善。首先，引入 OCR 技术识别批注图片，将图片笔记转化为可检索文本，提高知识利用率。其次，结合用户浏览、收藏、批注和学习路径行为，构建个性化书籍和资源推荐能力。再次，完善内容审核机制，引入敏感词过滤、举报处理和后台审核流程。最后，可探索与学校统一身份认证、课程信息或图书馆系统对接，使平台更贴合真实校园教学场景。")

    add_h1(doc, "参考文献")
    for idx, ref in enumerate(REFERENCES, 1):
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = Pt(17)
        para.paragraph_format.space_before = Pt(3)
        run = para.add_run(f"[{idx}] {ref}")
        apply_font(run, size=9)

    add_h1(doc, "致谢")
    add_para(doc, "本论文的完成离不开指导老师在选题、系统设计、论文结构和文档修改等方面给予的帮助。老师在论文撰写过程中多次指出存在的问题，使我能够及时调整技术描述和论文结构，在此表示衷心感谢。")
    add_para(doc, "感谢同学们在项目开发、功能测试和论文整理过程中提供的建议与帮助。感谢参考文献中各位作者的研究成果，为本文的需求分析和系统设计提供了有价值的参考。")
    add_para(doc, "最后，感谢家人和朋友在学习和毕业设计阶段给予的理解与支持。")

    OUT_DOC.parent.mkdir(exist_ok=True)
    doc.save(OUT_DOC)
    return OUT_DOC


if __name__ == "__main__":
    path = build_document()
    print(path.resolve())
