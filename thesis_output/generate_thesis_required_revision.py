# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "thesis_output"
DIAGRAM_DIR = OUT_DIR / "diagrams_required_revision"
OUT_DOC = OUT_DIR / "BookFlow_thesis_required_revision.docx"


REFERENCES = [
    "孙丽, 王皓, 戴璐, 等. 大学校园二手交易平台构建与运营——以E大学“花梨闲转”微信小程序为例[J]. 科技与创新, 2024(04): 12-16.",
    "高耀, 许诺, 李博, 等. 基于Web的新型校园二手交易平台实践研究——以“校易集市”为例[J]. 中国商论, 2023(01): 130-132.",
    "赵明, 贲祺舒, 吕怡. 基于微信小程序的高校校园二手物品交易平台的设计与实现[J]. 科技与创新, 2022(07): 55-58.",
    "沈政晔, 张辰瀚, 黄晋峰. 基于微信小程序的校园二手物品交易平台设计与开发[J]. 无线互联科技, 2021, 18(22): 66-68.",
    "王昱婷, 刘静, 燕明媚, 等. 基于微信小程序的大学生二手物品交易平台设计与开发[J]. 电脑知识与技术, 2019, 15(32): 283-284.",
    "曾倩. 基于微信小程序的高校二手物品交易和共享平台的搭建[J]. 现代商业, 2019(02): 52-53.",
    "彭嘉怡, 袁楠楠, 尹鲲龙, 等. 大学生闲置物品交易互助微信小程序设计与开发[J]. 中国新技术新产品, 2022(18): 37-39.",
    "沈莹, 黄旭, 曾孟佳. 基于SpringBoot+微信小程序的线上茶叶交易平台的设计与实现[J]. 福建茶叶, 2025(10): 49-51.",
    "李沛熹, 朱晓君, 姜建, 等. 基于微信小程序的校园二手书籍交易平台设计与实现[J]. 电脑知识与技术, 2021, 17(23): 57-59.",
    "陈怡婧, 郑晓溪, 李芳. 基于微信云开发的校园二手交易平台小程序的设计与实现[J]. 电脑知识与技术, 2022, 18(32): 51-54.",
]


INTERFACE_ROWS = [
    ("用户认证", "POST /user/auth/wechat", "完成微信小程序登录，返回用户信息和访问令牌。"),
    ("个人中心", "GET /user/profile", "查询用户资料、学生认证状态和个人基础信息。"),
    ("书籍管理", "GET /book/list；GET /book/detail", "完成书籍列表查询、详情展示和基础浏览。"),
    ("书籍发布", "POST /book/publish；POST /book/upload-image", "完成书籍信息发布与图片上传。"),
    ("订单交易", "POST /order/create；POST /order/pay/mock；POST /order/confirm-receipt", "完成订单创建、模拟支付和确认收货。"),
    ("批注传承", "GET /annotation/list；POST /annotation/create", "完成批注查询、文字批注和图片批注创建。"),
    ("资源管理", "POST /resource/upload-file；GET /resource/list", "完成资源上传、资源查询和资源绑定。"),
    ("学习路径", "GET /path/detail；POST /path/publish；POST /path/progress/complete-node", "完成学习路径展示、发布和进度更新。"),
    ("社区互动", "GET /community/feed；POST /community/post/create；POST /community/post/comment/create", "完成动态展示、帖子发布和评论互动。"),
    ("私信通知", "GET /chat/session/list；POST /chat/message/send；GET /user/notifications", "完成私信会话、消息发送和通知查看。"),
]


TEST_LOGIN_ROWS = [
    ("Test_01", "正常微信登录", "用户在登录页发起微信授权登录", "系统登录成功，保存会话状态并进入首页", "通过"),
    ("Test_02", "未登录访问受限页面", "未登录状态下进入个人中心或发布页", "系统提示先登录，并跳转到登录页面", "通过"),
    ("Test_03", "提交学生认证", "填写学号、姓名、学校等信息并提交认证材料", "系统保存认证信息，认证状态更新为待审核", "通过"),
    ("Test_04", "认证信息缺失", "缺少学号或姓名时提交认证申请", "系统阻止提交，并提示必填项不能为空", "通过"),
]


TEST_BOOK_TRADE_ROWS = [
    ("Test_01", "正常发布书籍", "上传封面图片并填写 ISBN、书名、作者、价格等信息后提交", "书籍发布成功，列表页可查看新书记录", "通过"),
    ("Test_02", "ISBN 自动补全", "输入系统支持识别的 ISBN 并触发自动补全", "系统自动回填书名、作者或出版社等书目信息", "通过"),
    ("Test_03", "创建订单", "买家在书籍详情页发起下单", "系统生成订单记录，订单状态更新为待支付", "通过"),
    ("Test_04", "订单状态流转", "买家完成模拟支付后，卖家发货，买家确认收货", "订单状态依次更新为待发货、待收货和已完成", "通过"),
    ("Test_05", "发布必填项校验", "未上传图片或未填写书名、价格时提交发布", "系统不允许提交，并给出对应校验提示", "通过"),
]


TEST_KNOWLEDGE_ROWS = [
    ("Test_01", "查看批注列表", "用户在书籍详情页进入批注列表页面", "系统正确显示该书籍下已有批注内容", "通过"),
    ("Test_02", "创建文字批注", "填写页码和批注内容后提交", "系统保存批注记录，并在列表页显示新增内容", "通过"),
    ("Test_03", "上传图片批注", "选择图片并填写批注说明后提交", "系统完成图片上传，批注内容可正常查看", "通过"),
    ("Test_04", "创建学习路径", "填写路径标题、难度和节点信息后发布", "系统保存学习路径，并在路径列表中可查询", "通过"),
    ("Test_05", "更新学习进度", "学习者进入路径详情并标记节点完成", "节点完成状态和整体学习进度同步更新", "通过"),
    ("Test_06", "查看节点资源", "在路径详情页点击节点资源入口", "系统显示与该节点绑定的学习资源列表", "通过"),
]


TEST_ADMIN_ROWS = [
    ("Test_01", "管理员登录", "管理员输入账号信息并进入后台", "系统登录成功，显示后台仪表盘与统计信息", "通过"),
    ("Test_02", "审核学生认证", "管理员查看认证列表并通过一条待审核记录", "认证状态更新成功，用户认证结果同步变化", "通过"),
    ("Test_03", "处理书籍状态", "管理员对违规书籍执行下架或删除操作", "系统完成状态更新，对应书籍不再正常展示", "通过"),
    ("Test_04", "审核社区内容", "管理员修改帖子或评论的展示状态", "系统保存审核结果，违规内容被限制显示", "通过"),
    ("Test_05", "处理举报与反馈", "管理员查看举报记录或反馈信息并执行处理", "系统记录处理结果，问题状态更新成功", "通过"),
]


DB_KEY_TABLES = [
    (
        "（1）用户基础表（wx_user）。该表用于保存用户的微信身份信息、昵称头像和最近登录状态，是系统登录与用户关联的基础表。",
        "表 4-1 用户基础表",
        [
            ("id", "bigint", "20", "用户主键", "是", "自增"),
            ("openid", "varchar", "64", "微信用户唯一标识", "否", "—"),
            ("unionid", "varchar", "64", "微信开放平台统一标识", "否", "—"),
            ("session_key", "varchar", "255", "微信会话密钥", "否", "—"),
            ("nickname", "varchar", "100", "用户昵称", "否", "—"),
            ("avatar_url", "varchar", "255", "头像地址", "否", "—"),
            ("mobile", "varchar", "20", "手机号", "否", "—"),
            ("last_login_time", "datetime", "—", "最后登录时间", "否", "—"),
        ],
    ),
    (
        "（2）用户认证表（user_profile）。该表用于保存学生认证资料和个人扩展信息，支撑学生认证审核和个人中心展示。",
        "表 4-2 用户认证表",
        [
            ("id", "bigint", "20", "认证记录主键", "是", "自增"),
            ("user_id", "bigint", "20", "关联 wx_user.id", "否", "—"),
            ("student_id", "varchar", "32", "学号", "否", "—"),
            ("real_name", "varchar", "50", "真实姓名", "否", "—"),
            ("school", "varchar", "100", "学校名称", "否", "—"),
            ("department", "varchar", "100", "院系信息", "否", "—"),
            ("auth_status", "int", "11", "认证状态", "否", "0"),
            ("credit_score", "int", "11", "信用分", "否", "80"),
            ("student_card_image_url", "varchar", "255", "学生证图片地址", "否", "—"),
            ("audit_remark", "varchar", "255", "审核备注", "否", "—"),
        ],
    ),
    (
        "（3）书籍信息表（book）。该表用于保存平台中的教材与二手书发布信息，是书籍浏览、检索和交易流程的核心业务表。",
        "表 4-3 书籍信息表",
        [
            ("id", "bigint", "20", "书籍主键", "是", "自增"),
            ("user_id", "bigint", "20", "发布者编号", "否", "—"),
            ("isbn", "varchar", "32", "书籍 ISBN 编号", "否", "—"),
            ("title", "varchar", "200", "书名", "否", "—"),
            ("author", "varchar", "100", "作者", "否", "—"),
            ("publisher", "varchar", "200", "出版社", "否", "—"),
            ("category", "varchar", "50", "书籍分类", "否", "—"),
            ("price", "decimal", "10,2", "售卖价格", "否", "—"),
            ("condition", "int", "11", "成色等级", "否", "—"),
            ("status", "int", "11", "书籍状态", "否", "1"),
            ("cover_images", "text", "—", "封面图片 JSON 数组", "否", "—"),
            ("description", "text", "—", "书籍描述", "否", "—"),
        ],
    ),
    (
        "（4）订单信息表（order）。该表用于保存一本书对应的一笔交易记录，支撑下单、支付、发货、收货和订单状态流转。",
        "表 4-4 订单信息表",
        [
            ("id", "bigint", "20", "订单主键", "是", "自增"),
            ("order_no", "varchar", "64", "订单编号", "否", "—"),
            ("book_id", "bigint", "20", "关联书籍编号", "否", "—"),
            ("buyer_id", "bigint", "20", "买家编号", "否", "—"),
            ("seller_id", "bigint", "20", "卖家编号", "否", "—"),
            ("total_amount", "decimal", "10,2", "订单金额", "否", "—"),
            ("status", "int", "11", "订单状态", "否", "0"),
            ("payment_method", "int", "11", "支付方式", "否", "—"),
            ("receiver_name", "varchar", "50", "收货人姓名", "否", "—"),
            ("receiver_phone", "varchar", "20", "收货人电话", "否", "—"),
            ("receiver_address", "varchar", "255", "收货地址", "否", "—"),
            ("payment_time", "datetime", "—", "支付时间", "否", "—"),
        ],
    ),
    (
        "（5）批注信息表（annotation）。该表用于保存书籍页码批注、图片批注和可见性信息，是学术资源传承模块的核心数据表。",
        "表 4-5 批注信息表",
        [
            ("id", "bigint", "20", "批注主键", "是", "自增"),
            ("book_id", "bigint", "20", "关联书籍编号", "否", "—"),
            ("user_id", "bigint", "20", "批注创建者编号", "否", "—"),
            ("page_num", "int", "11", "页码", "否", "—"),
            ("content", "text", "—", "文字批注内容", "否", "—"),
            ("position_text", "varchar", "255", "位置描述文本", "否", "—"),
            ("image_url", "varchar", "255", "批注图片地址", "否", "—"),
            ("type", "int", "11", "批注类型", "否", "1"),
            ("visibility", "int", "11", "可见范围", "否", "1"),
            ("like_count", "int", "11", "点赞数量", "否", "0"),
        ],
    ),
    (
        "（6）资源信息表（resource）。该表用于保存学习资料文件及其绑定关系，可将资源与书籍或学习路径节点关联起来。",
        "表 4-6 资源信息表",
        [
            ("id", "bigint", "20", "资源主键", "是", "自增"),
            ("user_id", "bigint", "20", "上传用户编号", "否", "—"),
            ("book_id", "bigint", "20", "关联书籍编号", "否", "—"),
            ("bind_type", "varchar", "32", "绑定类型", "否", "—"),
            ("bind_id", "bigint", "20", "绑定目标编号", "否", "—"),
            ("title", "varchar", "200", "资源标题", "否", "—"),
            ("type", "int", "11", "资源类别", "否", "—"),
            ("file_url", "varchar", "255", "文件访问地址", "否", "—"),
            ("file_size", "bigint", "20", "文件大小", "否", "—"),
            ("file_format", "varchar", "32", "文件格式", "否", "—"),
            ("visibility", "int", "11", "可见范围", "否", "1"),
        ],
    ),
    (
        "（7）学习路径表（learning_path）。该表用于保存用户创建的学习路线信息，描述路径标题、难度、状态和预估学习时长。",
        "表 4-7 学习路径表",
        [
            ("id", "bigint", "20", "路径主键", "是", "自增"),
            ("user_id", "bigint", "20", "创建者编号", "否", "—"),
            ("book_id", "bigint", "20", "关联书籍编号", "否", "—"),
            ("source_path_id", "bigint", "20", "来源路径编号", "否", "—"),
            ("title", "varchar", "200", "路径标题", "否", "—"),
            ("description", "text", "—", "路径描述", "否", "—"),
            ("cover_image", "varchar", "255", "封面图片", "否", "—"),
            ("difficulty", "int", "11", "难度等级", "否", "—"),
            ("estimated_hours", "int", "11", "预计学习时长", "否", "—"),
            ("status", "int", "11", "路径状态", "否", "0"),
        ],
    ),
    (
        "（8）路径节点表（path_node）。该表用于保存学习路径中的层次化节点，通过 parent_id 形成邻接表结构，并记录节点与资源的绑定信息。",
        "表 4-8 路径节点表",
        [
            ("id", "bigint", "20", "节点主键", "是", "自增"),
            ("path_id", "bigint", "20", "所属路径编号", "否", "—"),
            ("parent_id", "bigint", "20", "父节点编号", "否", "—"),
            ("title", "varchar", "200", "节点标题", "否", "—"),
            ("description", "text", "—", "节点描述", "否", "—"),
            ("order_num", "int", "11", "同级排序号", "否", "0"),
            ("estimated_minutes", "int", "11", "预计学习分钟数", "否", "—"),
            ("resource_ids", "text", "—", "关联资源编号列表", "否", "—"),
        ],
    ),
]


def font(size=24, bold=False):
    path = "C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc"
    if not Path(path).exists():
        path = "C:/Windows/Fonts/Deng.ttf"
    return ImageFont.truetype(path, size=size)


def wrap(text, chars):
    chars = int(chars)
    out = []
    for part in str(text).split("\n"):
        while len(part) > chars:
            out.append(part[:chars])
            part = part[chars:]
        out.append(part)
    return out


def draw_text_center(draw, box, text, size=22, bold=False):
    x1, y1, x2, y2 = box
    lines = wrap(text, max(2, (x2 - x1) // size))
    total = len(lines) * (size + 6)
    y = y1 + (y2 - y1 - total) / 2 + size / 2
    for line in lines:
        draw.text(((x1 + x2) / 2, y), line, fill="#111111", font=font(size, bold), anchor="mm")
        y += size + 6


def rect(draw, box, text, size=22, bold=True):
    draw.rectangle(box, fill="white", outline="#111111", width=3)
    draw_text_center(draw, box, text, size, bold)


def oval(draw, box, text, size=20):
    draw.ellipse(box, fill="white", outline="#111111", width=3)
    draw_text_center(draw, box, text, size, False)


def process(draw, box, text):
    draw.rounded_rectangle(box, radius=4, fill="white", outline="#111111", width=3)
    draw_text_center(draw, box, text, 22, True)


def terminator(draw, box, text):
    draw.rounded_rectangle(box, radius=35, fill="white", outline="#111111", width=3)
    draw_text_center(draw, box, text, 22, True)


def actor(draw, center_x, head_y, label):
    head_r = 28
    draw.ellipse((center_x - head_r, head_y, center_x + head_r, head_y + 2 * head_r), outline="#111111", width=3)
    neck_y = head_y + 2 * head_r
    body_bottom = neck_y + 120
    arm_y = neck_y + 45
    leg_y = body_bottom + 90
    draw.line([(center_x, neck_y), (center_x, body_bottom)], fill="#111111", width=3)
    draw.line([(center_x - 55, arm_y), (center_x + 55, arm_y)], fill="#111111", width=3)
    draw.line([(center_x, body_bottom), (center_x - 50, leg_y)], fill="#111111", width=3)
    draw.line([(center_x, body_bottom), (center_x + 50, leg_y)], fill="#111111", width=3)
    draw.text((center_x, leg_y + 45), label, fill="#111111", font=font(24, True), anchor="mm")
    return center_x + 55, arm_y


def usecase_boundary(draw, box, text):
    draw.rectangle(box, outline="#111111", width=4)
    x1, y1, _, _ = box
    draw.text((x1 + 20, y1 + 24), text, fill="#111111", font=font(24, True), anchor="la")


def decision(draw, cx, cy, w, h, text):
    pts = [(cx, cy - h / 2), (cx + w / 2, cy), (cx, cy + h / 2), (cx - w / 2, cy)]
    draw.polygon(pts, fill="white", outline="#111111")
    draw.line([pts[0], pts[1], pts[2], pts[3], pts[0]], fill="#111111", width=3)
    draw_text_center(draw, (cx - w / 3, cy - h / 4, cx + w / 3, cy + h / 4), text, 20, True)


def arrow(draw, start, end):
    draw.line([start, end], fill="#111111", width=3)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        d = 1 if ex > sx else -1
        pts = [(ex, ey), (ex - 14 * d, ey - 8), (ex - 14 * d, ey + 8)]
    else:
        d = 1 if ey > sy else -1
        pts = [(ex, ey), (ex - 8, ey - 14 * d), (ex + 8, ey - 14 * d)]
    draw.polygon(pts, fill="#111111")


def new_canvas(title, size=(1800, 1000)):
    img = Image.new("RGB", size, "white")
    draw = ImageDraw.Draw(img)
    draw.text((size[0] // 2, 55), title, fill="#111111", font=font(38, True), anchor="mm")
    return img, draw


def make_architecture():
    path = DIAGRAM_DIR / "04_01_role_architecture.png"
    img, d = new_canvas("图 4-1 系统总体架构图", (1900, 1050))
    rect(d, (690, 110, 1210, 180), "校园学术资源传承平台", 28)
    rect(d, (160, 270, 610, 340), "学生端微信小程序", 24)
    rect(d, (1280, 270, 1730, 340), "管理端后台系统", 24)
    for i, label in enumerate(["首页", "书籍", "交易", "批注", "学习路径", "社区", "个人中心"]):
        rect(d, (80 + i * 90, 420, 145 + i * 90, 620), label, 20)
    for i, label in enumerate(["仪表盘", "用户管理", "内容管理", "交易管理", "系统管理"]):
        rect(d, (1250 + i * 100, 420, 1320 + i * 100, 620), label, 20)
    rect(d, (550, 730, 1350, 800), "后端服务：用户、书籍、订单、批注、路径、资源、社区、私信", 22)
    rect(d, (360, 880, 710, 950), "数据库：核心业务数据持久化", 22)
    rect(d, (780, 880, 1120, 950), "缓存：登录态与临时数据", 22)
    rect(d, (1190, 880, 1540, 950), "文件服务：本地上传文件存储", 22)
    arrow(d, (950, 180), (385, 270))
    arrow(d, (950, 180), (1505, 270))
    arrow(d, (385, 340), (385, 420))
    arrow(d, (1505, 340), (1505, 420))
    arrow(d, (385, 620), (760, 730))
    arrow(d, (1505, 620), (1140, 730))
    arrow(d, (760, 800), (535, 880))
    arrow(d, (950, 800), (950, 880))
    arrow(d, (1140, 800), (1365, 880))
    img.save(path)
    return path


def make_module_tree():
    path = DIAGRAM_DIR / "04_02_function_module_tree.png"
    img, d = new_canvas("图 4-2 系统功能模块图", (1800, 1050))
    rect(d, (650, 110, 1150, 180), "校园学术资源传承平台", 28)
    rect(d, (230, 300, 580, 370), "用户端（小程序）", 26)
    rect(d, (1110, 300, 1460, 370), "管理端（Web）", 26)
    user_labels = ["首页模块", "书籍模块", "交易模块", "传承模块", "社区模块", "个人中心"]
    admin_labels = ["仪表盘", "用户管理", "内容管理", "交易管理", "系统管理"]
    ux = [120, 250, 380, 510, 640, 770]
    ax = [1080, 1210, 1340, 1470, 1600]
    for x, lab in zip(ux, user_labels):
        rect(d, (x, 520, x + 75, 740), lab, 22)
    for x, lab in zip(ax, admin_labels):
        rect(d, (x, 520, x + 75, 740), lab, 22)
    d.line([(900, 180), (900, 245), (405, 245), (405, 300)], fill="#111111", width=4)
    d.line([(900, 245), (1285, 245), (1285, 300)], fill="#111111", width=4)
    d.line([(405, 370), (405, 460)], fill="#111111", width=4)
    d.line([(157, 460), (807, 460)], fill="#111111", width=4)
    for x in [157, 287, 417, 547, 677, 807]:
        d.line([(x, 460), (x, 520)], fill="#111111", width=4)
    d.line([(1285, 370), (1285, 460)], fill="#111111", width=4)
    d.line([(1117, 460), (1637, 460)], fill="#111111", width=4)
    for x in [1117, 1247, 1377, 1507, 1637]:
        d.line([(x, 460), (x, 520)], fill="#111111", width=4)
    img.save(path)
    return path


def make_usecase(name, title, actor_label, cases):
    path = DIAGRAM_DIR / name
    img, d = new_canvas(title, (1900, 1150))
    boundary = (420, 140, 1750, 1020)
    usecase_boundary(d, boundary, "校园学术资源传承平台")
    actor_anchor = actor(d, 180, 390, actor_label)
    cols = [840, 1320]
    rows = (len(cases) + 1) // 2
    row_gap = 690 / max(rows - 1, 1) if rows > 1 else 0
    box_w, box_h = 340, 110
    start_y = 320
    for i, case in enumerate(cases):
        col = i % 2
        row = i // 2
        cx = cols[col]
        cy = int(start_y + row * row_gap)
        box = (cx - box_w // 2, cy - box_h // 2, cx + box_w // 2, cy + box_h // 2)
        oval(d, box, case, 21)
        d.line([actor_anchor, (box[0], cy)], fill="#111111", width=3)
    img.save(path)
    return path


def get_er_diagram():
    candidates = [
        OUT_DIR / "diagrams" / "core_ER_preview_clean.png",
        OUT_DIR / "figures" / "fig4_6_er.png",
        OUT_DIR / "diagrams" / "05_database_er_full.png",
    ]
    for path in candidates:
        if path.exists():
            return path
    raise FileNotFoundError("未找到可用的 ER 图文件")


def make_flow(name, title, steps):
    path = DIAGRAM_DIR / name
    img, d = new_canvas(title, (1900, 620))
    x, y = 80, 250
    terminator(d, (x, y, x + 190, y + 70), "开始")
    last = (x + 190, y + 35)
    x += 250
    for i, step in enumerate(steps):
        if step.startswith("判断"):
            decision(d, x + 100, y + 35, 200, 100, step)
            target = (x, y + 35)
            end = (x + 200, y + 35)
        else:
            process(d, (x, y, x + 220, y + 70), step)
            target = (x, y + 35)
            end = (x + 220, y + 35)
        arrow(d, last, target)
        last = end
        x += 280
    terminator(d, (x, y, x + 190, y + 70), "结束")
    arrow(d, last, (x, y + 35))
    img.save(path)
    return path


def generate_diagrams():
    DIAGRAM_DIR.mkdir(exist_ok=True)
    return {
        "user_usecase": make_usecase(
            "03_01_user_usecase.png",
            "图 3-1 用户用例图",
            "用户",
            [
                "微信登录与个人资料",
                "学生认证",
                "书籍浏览与搜索",
                "书籍发布",
                "订单交易",
                "批注传承",
                "学习路径与资源",
                "社区互动与私信",
            ],
        ),
        "admin_usecase": make_usecase(
            "03_02_admin_usecase.png",
            "图 3-2 管理员用例图",
            "管理员",
            [
                "后台登录",
                "用户与认证审核",
                "书籍与资源管理",
                "社区内容与举报处理",
                "订单与纠纷处理",
                "轮播图与通知维护",
            ],
        ),
        "er": get_er_diagram(),
        "architecture": make_architecture(),
        "module_tree": make_module_tree(),
        "login": make_flow("04_03_login_flow.png", "图 4-3 用户登录流程图", ["获取微信凭证", "提交登录接口", "判断用户是否存在", "生成Token", "返回登录结果"]),
        "order": make_flow("04_04_order_flow.png", "图 4-4 书籍交易流程图", ["浏览书籍", "创建订单", "模拟支付", "卖家发货", "买家收货", "更新订单状态"]),
        "annotation": make_flow("04_05_annotation_flow.png", "图 4-5 批注传承流程图", ["进入书籍详情", "查询批注", "填写批注", "上传图片", "保存批注", "展示批注"]),
        "path": make_flow("04_06_path_flow.png", "图 4-6 学习路径流程图", ["创建路径", "维护节点", "绑定资源", "发布路径", "开始学习", "更新进度"]),
        "test": make_flow("06_01_test_flow.png", "图 6-1 系统测试流程图", ["准备环境", "构造数据", "功能测试", "接口测试", "记录结果"]),
    }


def apply_font(run, name="宋体", size=10.5, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    for style_name, size in [("Heading 1", 18), ("Heading 2", 16), ("Heading 3", 14)]:
        st = doc.styles[style_name]
        st.font.name = "黑体"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        st.font.size = Pt(size)
        st.font.bold = True
    return doc


def p(doc, text="", first=True):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = Pt(20)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    if first:
        para.paragraph_format.first_line_indent = Cm(0.74)
    run = para.add_run(text)
    apply_font(run)
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style=None)
    para.paragraph_format.left_indent = Cm(0.74)
    para.paragraph_format.first_line_indent = Cm(-0.3)
    para.paragraph_format.line_spacing = Pt(20)
    run = para.add_run("（1）" + text)
    apply_font(run)
    return para


def h1(doc, text):
    doc.add_page_break()
    para = doc.add_paragraph(style="Heading 1")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing = Pt(36)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)
    run = para.add_run(text)
    apply_font(run, "黑体", 18, True)
    return para


def h2(doc, text):
    para = doc.add_paragraph(style="Heading 2")
    para.paragraph_format.line_spacing = Pt(24)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)
    run = para.add_run(text)
    apply_font(run, "黑体", 16, True)
    return para


def h3(doc, text):
    para = doc.add_paragraph(style="Heading 3")
    para.paragraph_format.line_spacing = Pt(22)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(text)
    apply_font(run, "黑体", 14, True)
    return para


def title(doc, text, size=18):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing = Pt(36)
    run = para.add_run(text)
    apply_font(run, "黑体", size, True)
    return para


def caption(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing = Pt(20)
    run = para.add_run(text)
    apply_font(run, "宋体", 9)


def figure(doc, path, cap, width=15.5):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(path), width=Cm(width))
    caption(doc, cap)


def set_cell(cell, text, bold=False):
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.line_spacing = Pt(17)
    run = para.add_run(str(text))
    apply_font(run, "宋体", 9, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, cap, headers, rows, widths=None):
    caption(doc, cap)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, head in enumerate(headers):
        set_cell(table.rows[0].cells[i], head, True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell(cells[i], val)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    return table


def add_toc(doc):
    doc.add_page_break()
    title(doc, "目录", 18)
    para = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-2" \\h \\z \\u')
    para._p.append(fld)


def build():
    diagrams = generate_diagrams()
    doc = setup_doc()
    title(doc, "基于微信小程序的校园学术资源传承平台的设计与实现", 18)
    title(doc, "摘要", 14)
    p(doc, "高校教材和专业书籍具有较强的课程属性，学生在课程结束或毕业离校后容易产生大量闲置书籍。传统线下转让方式依赖熟人关系，综合二手交易平台又存在场景不聚焦、检索成本较高等问题，难以充分满足校园内低单价、近距离、强时效的教材流转需求。同时，书籍中的批注、复习经验和配套资料往往随书籍闲置而散失，无法持续服务后续学习者。")
    p(doc, "针对上述问题，本文设计并实现了一个基于微信小程序的校园学术资源传承平台。系统采用前后端分离架构，学生端基于 UniApp 构建微信小程序页面，管理端基于 Vue 3 构建 Web 后台，后端基于 Spring Boot 提供业务接口，并使用 MySQL 保存核心业务数据。平台围绕“书籍流转”和“知识传承”两条主线，提供用户认证、学生认证、书籍发布、订单交易、批注传承、学习路径、资源上传、社区互动和私信通知等功能。")
    p(doc, "本文按照软件工程方法完成需求分析、总体设计、数据库设计、功能实现和系统测试。数据库设计以用户、书籍、订单、批注、资源、学习路径等实体为核心，能够支撑平台的主要业务流程。测试结果表明，系统能够完成核心功能闭环，具有一定的实用价值和扩展空间。")
    p(doc, "关键词：校园二手书    学术资源传承    UniApp    Spring Boot    MySQL", False)
    add_toc(doc)

    h1(doc, "1 绪论")
    h2(doc, "1.1 研究背景")
    p(doc, "随着互联网应用和电子商务的发展，高校学生的线上消费场景不断增加，教材、教辅、电子产品和生活用品等闲置物品也随之增多。孙丽等在对 E 大学校园二手交易平台的研究中指出，校园二手交易具有较大发展空间，但仍存在规范性不足、运营机制不完善等问题[1]。高耀等在“校易集市”相关研究中也提到，校园内部缺少统一交易平台时，学生往往依赖微信群、QQ群或线下跳蚤市场发布信息，容易出现信息覆盖、分类检索困难和交易效率不高等情况[2]。")
    p(doc, "对于高校学生而言，二手交易不仅是闲置物品再利用方式，也是降低学习和生活成本的重要途径。赵明等认为，校园二手交易平台能够为学生提供长期、稳定的交易渠道，并在同校交易场景下增强交易便利性和安全性[3]。沈政晔等从共享经济和校园安全角度分析了微信小程序在校园二手交易中的适用性，指出小程序可以围绕登录、首页、发布和个人中心等功能组织平台业务[4]。")
    p(doc, "在校园二手交易中，书籍具有较强的特殊性。高校学生在课程学习中会购买大量教材和参考书，课程结束后部分书籍处于闲置状态，但书中形成的批注、复习重点和资料整理仍具有学习价值。现有校园二手平台多关注商品发布、搜索、沟通和订单处理，对书籍中的学习经验传递关注不足。因此，有必要在二手书交易基础上进一步引入批注传承、资源共享和学习路径组织等功能。")
    h2(doc, "1.2 研究目的与意义")
    p(doc, "本文旨在设计并实现一个面向高校学生的校园学术资源传承平台，使二手书交易与批注传承、资源共享和学习路径组织结合起来。平台不仅提供书籍发布、浏览、沟通和交易等基础功能，还将书籍中的批注信息、学习资料和路径节点进行数字化管理，使闲置教材在流转过程中保留更多学习价值。")
    p(doc, "从实际应用角度看，该系统有三方面意义。第一，平台能够提升校园二手书流转效率，减少学生寻找买家或卖家的时间成本。第二，平台通过批注、资源和学习路径功能，将分散在书本和个人笔记中的学习经验沉淀为可查询、可复用的数字内容。第三，平台引入学生认证、个人中心、后台审核等机制，有助于提升校园交易场景下的身份可信度和内容规范性。王昱婷等提出，大学生二手物品交易平台需要面向学生群体提升操作便利性、推荐个性化和信息对等性[5]。")
    p(doc, "曾倩在高校二手物品交易和共享平台研究中指出，书籍在高校闲置物品供需中占比较高，具备在校内实现供需匹配和循环利用的条件[6]。彭嘉怡等在校园闲置物品交易互助小程序研究中强调，校园认证、个性化推荐和消息提醒能够提升平台的安全性和服务效率[7]。本系统在此基础上进一步聚焦二手书籍和学习资源传承场景，使平台既服务于教材流转，也服务于学习经验沉淀。")
    p(doc, "从工程实践角度看，本系统采用 Spring Boot、UniApp、Vue 3 和 MySQL 等技术完成前后端分离开发，覆盖微信小程序页面设计、后台管理端设计、REST 接口设计、关系数据库建模、文件上传、鉴权和系统测试等内容。沈莹等在基于 SpringBoot 与微信小程序的交易平台研究中说明，后端模块化设计与小程序前端结合能够支撑用户管理、商品展示和订单处理等业务功能[8]。因此，本系统具有较明确的软件工程实践价值。")
    h2(doc, "1.3 国内外研究现状")
    p(doc, "现有研究主要从校园闲置物品流通、平台运营、微信小程序开发和交易安全等角度展开。相关研究普遍认为，校园二手平台应围绕学生真实交易场景进行设计，重点解决信息发布分散、分类检索困难、买卖双方沟通不便和交易规范性不足等问题。与综合性二手交易平台相比，校园平台的优势在于服务对象集中、交易距离较短、身份关系较清晰，能够更好地满足高校学生低单价、高频次、近距离的交易需求。")
    p(doc, "针对书籍交易场景，李沛熹等认为校园二手书籍交易平台能够降低学生学习成本，并推动低碳可持续发展[9]。不过，现有校园二手书平台通常仍以书籍买卖为主要目标，对书籍使用过程中形成的批注、复习重点、学习资料和课程路径缺少系统化沉淀。本文在已有研究基础上，将二手书交易、批注传承、学习资源和学习路径组织放在同一平台中进行设计，重点解决书籍流转与学习经验传递之间的衔接问题。")
    p(doc, "在微信小程序实现方式方面，陈怡婧等基于微信云开发实现了校园二手交易平台小程序，并将留言、回复、商品分类、关键词查找和附近高校选择作为特色功能[10]。这说明微信小程序能够较好地承载校园二手交易中的移动端浏览、发布、搜索和沟通需求。结合本系统的业务目标，本文在用户端采用微信小程序形式，在后台提供管理端功能，使用户操作和平台维护能够形成完整闭环。")
    h2(doc, "1.4 本文主要工作")
    p(doc, "本文主要完成以下工作：分析平台业务需求，明确用户和管理员的功能边界；设计系统总体架构和标准功能模块图；完成用户、书籍、订单、批注、资源、学习路径等核心数据模型设计；基于 Spring Boot、UniApp、MySQL 等技术实现主要功能；最后对核心业务流程进行功能和接口测试。")
    h2(doc, "1.5 论文组织结构")
    p(doc, "全文共分七章：第一章为绪论；第二章介绍主要技术选型及原因；第三章进行系统需求分析；第四章完成系统总体设计；第五章按功能模块说明系统实现；第六章进行系统测试；第七章总结工作并提出展望。")

    h1(doc, "2 相关技术综述")
    h2(doc, "2.1 Spring Boot 框架")
    p(doc, "Spring Boot 是基于 Spring 生态体系的 Java 后端开发框架，适合用于构建中小型 Web 应用和 REST 风格接口服务。与传统 Spring 项目相比，Spring Boot 通过自动配置、内置 Web 容器和统一依赖管理简化了项目搭建过程，能够减少大量重复配置工作，使开发者将主要精力放在业务功能实现上。对于本系统而言，后端需要同时支撑用户认证、书籍发布、订单交易、批注管理、学习路径、资源上传、社区互动和私信沟通等多个模块，业务接口数量较多，采用 Spring Boot 可以较好地组织 Controller、Service、Mapper 等层次结构。")
    p(doc, "本系统选择 Spring Boot 的主要原因包括三点。第一，Spring Boot 对 Web 接口开发支持成熟，能够快速提供 REST API，适合前后端分离架构下的小程序端和 Web 管理端调用。第二，Spring Boot 生态完善，能够与数据库访问、文件上传、跨域配置、拦截器和统一异常处理等功能自然集成，便于实现完整业务闭环。第三，Spring Boot 项目结构清晰，适合毕业设计中展示软件工程的分层思想。在本系统中，Spring Boot 主要承担业务处理和接口服务作用，是连接前端页面、数据库和文件服务的核心后端框架。")
    h2(doc, "2.2 UniApp 与微信小程序")
    p(doc, "UniApp 是一套使用 Vue 语法进行跨端开发的前端框架，能够将同一套代码编译到微信小程序、H5、App 等不同平台。本系统的主要使用对象是高校学生，学生日常使用移动端频率较高，而微信小程序具有无需安装、打开方便、分享便捷等特点，因此用户端选择微信小程序作为主要承载形式。使用 UniApp 可以在保持微信小程序使用体验的同时，降低页面开发和维护成本。")
    p(doc, "本系统选择 UniApp 的原因主要体现在三个方面。第一，UniApp 使用组件化和数据驱动方式开发页面，适合构建首页、分类、发布、书籍详情、批注、学习路径和个人中心等多页面应用。第二，UniApp 对微信小程序生态支持较好，可以调用登录、上传、页面跳转等小程序能力，满足用户认证和文件上传等业务需要。第三，UniApp 语法与 Vue 体系接近，便于与管理端的 Vue 技术栈保持开发习惯一致。在本系统中，UniApp 主要负责学生端页面展示、用户操作采集和后端接口调用，是学生用户使用平台的主要入口。")
    h2(doc, "2.3 Vue 3 与管理后台")
    p(doc, "Vue 3 是当前常用的前端框架之一，采用组件化、响应式数据绑定和组合式 API 等机制，适合构建交互较多的管理后台页面。校园学术资源传承平台除了学生端小程序外，还需要提供后台管理端，用于辅助完成用户认证审核、内容管理、订单问题处理和系统运营配置等工作。此类页面通常包含数据表格、筛选条件、表单弹窗和状态操作，使用 Vue 3 能够较好地组织这些界面逻辑。")
    p(doc, "本系统选择 Vue 3 作为管理端技术的原因主要包括三点。第一，Vue 3 组件化能力较强，可以将用户列表、书籍管理、订单管理、内容审核等页面拆分为独立组件，提升前端代码复用性。第二，Vue 3 与 Element Plus 等后台管理生态结合成熟，能够快速构建表格、表单、菜单和布局等常见后台界面。第三，管理端主要面向管理员使用，对数据展示和操作效率要求较高，Vue 3 的响应式更新机制能够提高页面交互体验。在本系统中，Vue 3 管理后台主要承担后台数据查看、内容维护、认证审核和运营辅助等作用。")
    h2(doc, "2.4 MySQL 数据库")
    p(doc, "MySQL 是本系统最核心的数据存储技术之一。平台中的用户、书籍、订单、批注、资源、学习路径、社区帖子和私信消息都具有明确的结构化字段和实体关系，适合使用关系型数据库进行管理。与非关系型数据库相比，MySQL 在事务处理、表结构约束、条件查询和数据一致性方面更适合本系统的核心业务。")
    p(doc, "本系统选择 MySQL 的原因主要包括三点。第一，书籍交易、订单状态和用户信息需要稳定持久化，MySQL 能够提供可靠的数据存储能力，避免核心业务数据丢失。第二，系统实体之间存在清晰的一对多或多对一关系，例如用户与书籍、书籍与批注、订单与书籍、学习路径与路径节点，关系型数据库更便于表达和维护这些关联。第三，MySQL 查询能力成熟，能够支持书籍列表、分类检索、订单查询、个人中心统计和后台管理查询等业务场景。")
    p(doc, "在系统中的作用上，MySQL 主要承担业务数据持久化任务。用户基础信息、学生认证资料、书籍数据、订单状态、批注内容、学习资源、学习路径节点、社区帖子和私信消息均保存在 MySQL 中。通过这些数据表的组合，系统能够支撑书籍流转、知识传承、学习资源管理和社区互动等核心功能。因此，MySQL 不仅是数据存储工具，也是系统业务关系和数据一致性的基础。")

    h1(doc, "3 系统需求分析")
    h2(doc, "3.1 系统可行性分析")
    p(doc, "从技术可行性看，Spring Boot、UniApp、Vue 3 和 MySQL 均为成熟技术，能够支撑本系统的页面交互、接口服务、管理后台和数据持久化。从经济可行性看，系统可在普通开发环境下部署测试，不依赖高成本商业服务。从操作可行性看，微信小程序符合学生移动端使用习惯，后台管理端也便于管理员进行内容和订单辅助处理。")
    h2(doc, "3.2 用户角色分析")
    p(doc, "参照系统开发类论文中常见的角色划分方式，本系统将使用者划分为用户和管理员两类。用户是平台前台业务的主要参与者，负责完成书籍浏览、发布、交易、批注、学习路径、资源共享和社区互动等操作；管理员负责后台审核与维护工作，主要处理用户认证、内容管理、订单问题和基础运营维护。为明确不同角色的业务边界，本文分别给出用户用例分析和管理员用例分析。")
    h3(doc, "3.2.1 用户用例分析")
    p(doc, "用户主要通过微信小程序访问系统，能够完成微信登录、学生认证、书籍浏览与搜索、书籍发布、订单交易、批注传承、学习路径与资源管理、社区互动和私信沟通等业务操作。用户用例如图 3-1 所示。")
    figure(doc, diagrams["user_usecase"], "图 3-1 用户用例图", 16)
    h3(doc, "3.2.2 管理员用例分析")
    p(doc, "管理员主要通过 Web 管理端访问系统，能够完成后台登录、用户与认证审核、书籍与资源管理、社区内容与举报处理、订单与纠纷处理以及轮播图与通知维护等操作。管理员用例如图 3-2 所示。")
    figure(doc, diagrams["admin_usecase"], "图 3-2 管理员用例图", 16)
    h2(doc, "3.3 功能需求分析")
    p(doc, "根据使用者角色的不同，系统功能需求采用列表形式描述如下。")
    p(doc, "用户功能需求：", False)
    for item in [
        "账号与认证功能。用户可以通过微信小程序登录系统，查看和维护个人资料，提交学号、姓名和认证材料完成学生认证。",
        "书籍浏览与搜索功能。用户可以在首页查看推荐书籍，通过分类、关键词和详情页了解书籍信息。",
        "书籍发布功能。完成学生认证后的用户可以发布二手书信息，填写 ISBN、书名、作者、价格、成色和描述，并上传书籍图片。",
        "订单交易功能。购买者可以基于书籍创建订单，完成模拟支付、查看订单详情、确认收货；卖家可以处理发货和订单问题。",
        "批注传承功能。用户可以查看书籍批注，创建文字批注或图片批注，并对有价值的批注进行点赞。",
        "学习路径与资源功能。用户可以创建学习路径、维护路径节点、上传资源文件，并记录学习进度。",
        "社区与私信功能。用户可以发布社区动态、评论帖子、举报内容，并通过私信与交易对象沟通。"]:
        bullet(doc, item)
    p(doc, "管理员功能需求：", False)
    for item in [
        "用户管理功能。管理员可以查看用户信息，对学生认证材料进行审核和辅助管理。",
        "内容管理功能。管理员可以管理书籍、批注、资源、社区动态、评论和举报记录，维护平台内容质量。",
        "交易管理功能。管理员可以查看订单信息，处理订单问题和纠纷记录。",
        "系统运营功能。管理员可以维护轮播图、查看通知与反馈，并进行基础数据概览。"]:
        bullet(doc, item)
    h2(doc, "3.4 非功能需求分析")
    p(doc, "系统非功能需求包括易用性、安全性、可维护性、兼容性和数据一致性。易用性要求小程序页面流程清晰，发布、搜索和下单等操作符合学生使用习惯；安全性要求登录后接口通过 token 校验；可维护性要求后端按 Controller、Service、Mapper 和 Entity 分层；兼容性要求小程序端适配微信运行环境；数据一致性要求订单、批注和学习进度等状态由后端统一维护。")

    h1(doc, "4 系统总体设计")
    h2(doc, "4.1 系统总体架构设计")
    p(doc, "系统采用前后端分离架构，并在架构图中将学生端和管理端的功能模块分开绘制。学生端由微信小程序承载，主要面向学生用户；管理端由 Web 后台承载，主要面向管理员。二者通过 HTTP/HTTPS 与 Spring Boot 后端通信，后端再访问 MySQL、Redis 和本地文件存储服务。系统总体架构如图 4-1 所示。")
    figure(doc, diagrams["architecture"], "图 4-1 系统总体架构图", 16)
    h2(doc, "4.2 系统功能模块设计")
    p(doc, "系统功能模块图采用软件工程中自顶向下的层次化分解方式绘制。顶层为校园学术资源传承平台，一级模块分为用户端（小程序）和管理端（Web）。用户端进一步分解为首页模块、书籍模块、交易模块、传承模块、社区模块和个人中心；管理端进一步分解为仪表盘、用户管理、内容管理、交易管理和系统管理。系统功能模块如图 4-2 所示。")
    figure(doc, diagrams["module_tree"], "图 4-2 系统功能模块图", 16)
    h2(doc, "4.3 业务流程设计")
    p(doc, "为保证图表清晰，本文业务流程均采用标准流程图样式进行绘制。用户登录、书籍交易、批注传承和学习路径流程分别如图 4-3 至图 4-6 所示。")
    figure(doc, diagrams["login"], "图 4-3 用户登录流程图", 16)
    figure(doc, diagrams["order"], "图 4-4 书籍交易流程图", 16)
    figure(doc, diagrams["annotation"], "图 4-5 批注传承流程图", 16)
    figure(doc, diagrams["path"], "图 4-6 学习路径流程图", 16)
    h2(doc, "4.4 数据库设计")
    p(doc, "数据库设计按照概念数据模型设计和物理模型设计两个层次展开。概念数据模型用于描述系统中主要实体及其联系，物理模型设计则进一步落实为具体的数据表、字段和约束结构。")
    h3(doc, "4.4.1 概念数据模型设计")
    p(doc, "概念数据模型设计采用 E-R 图对系统核心实体及实体关系进行抽象描述。从业务角度看，平台中的主要实体包括用户、学生认证、书籍、订单、批注、资源、学习路径、路径节点、社区帖子、评论、私信会话和私信消息等对象。各实体之间通过发布、认证、交易、创建、评论和会话等关系构成完整的数据联系，其中用户与学生认证是一对一关系，用户与书籍、学习路径、帖子等实体是一对多关系，学习路径与路径节点、帖子与评论、会话与消息也分别形成一对多关系。系统概念数据模型如图 4-7 所示。")
    figure(doc, diagrams["er"], "图 4-7 数据库 E-R 图", 16)
    h3(doc, "4.4.2 物理模型设计")
    p(doc, "物理模型设计是在概念模型基础上，将实体转换为实际数据库表结构，并明确字段名称、字段类型、长度、主键和默认值等内容。参考系统开发类论文中物理模型结构设计的写法，本文从平台核心业务中选取用户、认证、书籍、订单、批注、资源、学习路径和路径节点 8 张关键表进行说明。")
    for intro_text, table_caption, table_rows in DB_KEY_TABLES:
        p(doc, intro_text, False)
        add_table(
            doc,
            table_caption,
            ["字段名称", "类型", "长度", "字段说明", "主键", "默认值"],
            table_rows,
            [2.3, 2.5, 1.4, 6.2, 1.2, 2.1],
        )
    p(doc, "除上述关键表外，post、comment、chat_session、chat_message、content_report 和 notification 等表分别用于支撑社区互动、私信沟通、内容举报和系统通知等辅助业务。")
    h2(doc, "4.5 接口设计")
    p(doc, "系统接口设计遵循 REST 风格，前后端通过 HTTP/HTTPS 协议和 JSON 数据格式完成交互，文件上传接口采用 multipart/form-data 方式实现。为避免正文接口表过长，本文仅选取与系统核心业务直接相关的代表性接口进行说明。公开接口主要用于书籍浏览、路径详情和社区动态展示，涉及用户资料、订单处理、批注创建、资源上传和私信发送等操作的接口需要在请求头中携带 Authorization 令牌。核心接口设计如表 4-9 所示。")
    add_table(doc, "表 4-9 核心接口设计表", ["功能模块", "代表接口", "功能说明"], INTERFACE_ROWS, [3.0, 7.2, 6.2])

    h1(doc, "5 系统详细实现")
    p(doc, "系统详细实现部分参考系统开发类论文的写法，重点说明各功能模块面向用户和管理员提供了哪些具体操作能力，以及这些功能如何支撑校园学术资源传承平台的完整业务闭环，而不再单纯按照接口调用顺序展开描述。")
    h2(doc, "5.1 用户端功能实现")
    h3(doc, "5.1.1 首页浏览与分类检索")
    p(doc, "用户进入小程序后，首先可以在首页查看平台轮播图、热门书籍、推荐学习路径和社区动态等内容。首页将书籍浏览、路径学习、书籍发布和社区交流等入口集中展示，便于用户快速进入核心业务页面。对于有明确目标的用户，系统还提供分类页和搜索页，支持按照书籍分类、关键词等方式检索目标资源，从而提高二手书和学习资料的查找效率。")
    h3(doc, "5.1.2 登录认证与个人中心")
    p(doc, "当用户首次使用系统时，可以通过微信授权完成登录，系统在登录成功后保存当前会话状态，并在后续访问需要身份校验的页面时自动进行登录检查。进入个人中心后，用户可以完善昵称、头像、院系、简介等个人资料，也可以提交学生认证信息。除身份信息维护外，个人中心还整合了我的书架、我的订单、我的收藏、收货地址、通知中心、我的批注、我创建的路径和我上传的资源等入口，使用户能够集中管理个人资产和学习成果。")
    h3(doc, "5.1.3 书籍发布与订单交易")
    p(doc, "在书籍发布功能中，用户可以上传书籍图片，填写 ISBN、书名、作者、出版社、分类、成色、价格和描述等信息，并根据实际需要设置批注可见范围和关联学习路径。系统支持通过 ISBN 自动补全部分书目信息，减少手工录入工作量。书籍发布成功后，其他用户可在详情页查看封面、价格、书况和附加信息，并进一步发起沟通或下单。订单交易功能围绕一本书对应一笔订单的模式展开，用户端支持下单、模拟支付、查看订单状态、卖家发货和买家确认收货等操作，从而形成较为完整的校园二手书交易流程。")
    h3(doc, "5.1.4 批注传承与资源查看")
    p(doc, "批注传承是本系统区别于普通二手书交易平台的重要功能之一。用户在浏览书籍详情时，不仅可以了解书籍的基本信息，还可以进入批注列表页查看前任使用者留下的学习记录。在具备权限的情况下，用户可以按页码新增文字批注或图片批注，并通过点赞等方式对优质内容进行反馈。与此同时，系统还提供资源查看功能，支持用户围绕书籍或学习路径节点查看和使用相关学习资料，使书籍交易不再停留在纸质资源流转层面，而是进一步延伸到附加知识内容的传承。")
    h3(doc, "5.1.5 学习路径与进度管理")
    p(doc, "对于希望系统化学习某一课程或方向的用户，平台提供了学习路径功能。用户可以查看路径标题、难度、创建者、预计学习时长和节点说明，并在学习过程中按节点逐步推进。系统以节点清单的形式展示学习顺序，用户可以查看节点对应的学习资源、标记完成状态，并根据已完成节点数量实时查看学习进度。对于路径创建者，系统支持新建学习路径和维护路径节点；对于普通学习者，则重点支持路径收藏、开始学习和持续跟进，增强平台的学习组织能力。")
    h3(doc, "5.1.6 社区互动与私信沟通")
    p(doc, "社区模块为用户提供了书评分享、学习问答、路径交流和动态浏览等互动场景。用户可以发布帖子，查看其他用户发布的经验内容，并对帖子执行点赞、评论、收藏和举报等操作。除公开交流外，系统还提供私信沟通功能，尤其适用于买卖双方围绕具体书籍进行交易协商。聊天页面除消息发送外，还同步展示当前书籍和订单的关键信息，并根据交易阶段提供下单、支付、取消订单、发货和确认收货等快捷操作，从而将沟通行为与交易流程有效结合。")
    h2(doc, "5.2 管理端功能实现")
    h3(doc, "5.2.1 后台登录与数据概览")
    p(doc, "管理员通过后台管理端登录系统后，可以进入数据概览页面查看平台的总体运行情况。后台支持展示核心统计信息和日志信息，便于管理员了解当前用户规模、内容数量、交易情况及系统运行状态。该功能为后续审核、维护和运营管理提供了统一的工作入口。")
    h3(doc, "5.2.2 用户与认证审核管理")
    p(doc, "在用户管理方面，管理员可以按关键词、状态和认证状态筛选用户信息，并对异常账号进行状态调整。对于学生认证业务，后台提供专门的认证审核功能，管理员可以查看用户提交的认证资料并给出审核结果。通过这一模块，平台能够在保证用户准入规范性的同时，维持较为可靠的校园身份环境。")
    h3(doc, "5.2.3 书籍、资源与学习路径管理")
    p(doc, "针对平台核心内容资源，后台管理端提供了书籍、资源和学习路径的统一管理能力。管理员可以查看书籍发布信息，对不符合规范的内容进行状态处理或删除；对于用户上传的资源和创建的学习路径，也可以按状态进行审核和维护。这样既能够保证资源内容的可用性和规范性，也有助于提升平台整体内容质量。")
    h3(doc, "5.2.4 订单、纠纷与社区内容管理")
    p(doc, "在交易管理方面，后台支持查看订单列表和纠纷信息，便于管理员掌握交易执行情况并介入异常订单处理。在社区内容管理方面，管理员可以分别查看帖子和评论，对违规或不适宜展示的内容进行状态调整。该部分功能使平台既能支持用户自由交流，又能通过后台治理机制控制内容风险和交易风险。")
    h3(doc, "5.2.5 轮播图、举报与反馈处理")
    p(doc, "除核心业务管理外，后台还提供轮播图维护、内容举报处理和用户反馈处理等辅助功能。管理员可以维护首页展示内容，提升平台首页的信息组织效果；对于用户提交的举报和意见反馈，后台能够进行统一查看和处理，形成从问题发现、人工审核到结果反馈的闭环。这些功能虽然不直接参与交易和学习流程，但对平台的日常运营和持续优化具有重要作用。")

    h1(doc, "6 系统测试")
    h2(doc, "6.1 测试目的")
    p(doc, "系统测试是系统开发过程中的重要环节，其主要目的是验证校园学术资源传承平台是否能够按照需求稳定运行，检查用户端和管理端的主要功能是否满足设计要求，并确认关键业务流程在实际操作过程中能够得到正确反馈。通过测试，还可以及时发现功能实现中的异常情况，降低系统投入使用后出现错误的概率，从而保证平台具有较好的可用性和稳定性。")
    h2(doc, "6.2 测试方法")
    p(doc, "本系统测试主要采用黑盒测试方法。黑盒测试不关注系统内部代码实现细节，而是从用户操作和功能结果出发，根据需求分析和功能设计内容编写测试用例，检验系统输出是否与预期一致。在实际测试过程中，首先在本地开发环境中部署后端服务、数据库和前端页面，然后围绕用户登录、书籍发布、订单交易、批注传承、学习路径和后台管理等主要功能逐项执行测试；当发现问题后，对相关功能进行修正，并再次执行回归验证，确保修改不会影响其他模块的正常运行。")
    h2(doc, "6.3 测试内容")
    p(doc, "结合本系统的功能结构，测试内容重点围绕用户登录与认证、书籍发布与交易、批注传承与学习路径、后台管理等主要功能展开。测试时针对不同模块分别设计测试用例，并通过实际操作验证系统响应是否正确。")
    p(doc, "（1）用户登录与认证模块测试。用户登录与认证测试用例表如表 6-1 所示。", False)
    add_table(doc, "表 6-1 用户登录与认证测试用例表", ["用例序号", "测试用例", "输入数据或操作", "系统返回数据", "测试结果"], TEST_LOGIN_ROWS, [2.2, 3.0, 5.6, 5.2, 1.8])
    p(doc, "（2）书籍发布与交易模块测试。书籍发布与交易测试用例表如表 6-2 所示。", False)
    add_table(doc, "表 6-2 书籍发布与交易测试用例表", ["用例序号", "测试用例", "输入数据或操作", "系统返回数据", "测试结果"], TEST_BOOK_TRADE_ROWS, [2.2, 3.0, 5.8, 5.0, 1.8])
    p(doc, "（3）批注传承与学习路径模块测试。批注传承与学习路径测试用例表如表 6-3 所示。", False)
    add_table(doc, "表 6-3 批注传承与学习路径测试用例表", ["用例序号", "测试用例", "输入数据或操作", "系统返回数据", "测试结果"], TEST_KNOWLEDGE_ROWS, [2.2, 3.2, 5.6, 5.0, 1.8])
    p(doc, "（4）后台管理模块测试。后台管理测试用例表如表 6-4 所示。", False)
    add_table(doc, "表 6-4 后台管理测试用例表", ["用例序号", "测试用例", "输入数据或操作", "系统返回数据", "测试结果"], TEST_ADMIN_ROWS, [2.2, 3.0, 5.6, 5.2, 1.8])
    h2(doc, "6.4 测试结论")
    p(doc, "经过对系统主要功能模块的测试可以看出，用户端登录认证、资料维护、书籍发布、订单交易、批注创建、学习路径学习、资源上传、社区互动以及管理端审核维护等功能均能够按照预期完成，测试用例结果均为通过，系统能够对用户操作给出正确反馈，主要业务流程运行正常。总体来看，系统已经具备较为完整的功能，能够满足校园学术资源传承平台的基本使用需求，并可在后续应用过程中继续进行维护、优化和功能升级。")

    h1(doc, "7 总结与展望")
    h2(doc, "7.1 工作总结")
    p(doc, "本文完成了校园学术资源传承平台的需求分析、系统设计、数据库设计、功能实现和测试验证。系统将二手书流转与批注传承、学习路径和资源共享结合起来，能够为高校学生提供教材交易和学习经验沉淀服务。")
    h2(doc, "7.2 系统不足")
    p(doc, "当前系统仍存在不足：内容审核自动化程度不高，批注图片尚未进行 OCR 识别，推荐功能仍较基础，支付流程采用模拟方式，系统也尚未进行大规模压力测试。")
    h2(doc, "7.3 未来展望")
    p(doc, "后续可从智能推荐、OCR 批注识别、统一身份认证、内容审核和性能优化等方面继续完善系统，使平台更加适合真实校园环境中的长期使用。")

    h1(doc, "参考文献")
    for i, ref in enumerate(REFERENCES, 1):
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = Pt(17)
        run = para.add_run(f"[{i}] {ref}")
        apply_font(run, "宋体", 9)
    h1(doc, "致谢")
    p(doc, "本论文的完成离不开指导老师在选题、系统设计和论文修改过程中给予的帮助。老师多次指出论文结构、图表和技术描述中的问题，使本文能够逐步完善，在此表示衷心感谢。")
    p(doc, "感谢同学们在项目开发和功能测试过程中提供的建议，也感谢家人和朋友在毕业设计阶段给予的理解与支持。")
    doc.save(OUT_DOC)
    return OUT_DOC


if __name__ == "__main__":
    print(build().resolve())
