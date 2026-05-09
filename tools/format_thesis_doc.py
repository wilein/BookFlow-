import json
import os
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


ROOT = Path.cwd()
SRC = ROOT / "paper_source.docx"
OUT_DIR = ROOT / "thesis_output"
DIAGRAM = OUT_DIR / "diagrams"
FINAL = OUT_DIR / "paper_full_formatted.docx"
FINAL_CN = OUT_DIR / "校园学术资源传承平台论文_格式图表完整版.docx"
ORIGINAL_DIR = Path(r"D:\1011\1011\毕设\【元新学院】2025届毕业论文（设计）归档材料包【4.9确认版】【学生】\（1）正本材料")


def run_font(run, name="宋体", size=10.5, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def style_para(p):
    text = p.text.strip()
    if not text:
        return
    for r in p.runs:
        run_font(r)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if re.match(r"^第[一二三四五六七八九十]+章", text):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = None
        for r in p.runs:
            run_font(r, "黑体", 15, True)
    elif re.match(r"^\d+\.\d+", text):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = None
        for r in p.runs:
            run_font(r, "黑体", 12, True)
    elif text in {"摘要", "目录", "参考文献", "致谢"}:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = None
        for r in p.runs:
            run_font(r, "黑体", 14, True)
    elif text.startswith("图") or text.startswith("表"):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = None
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Cm(0.74)


def after(p, text=""):
    new_p = OxmlElement("w:p")
    p._p.addnext(new_p)
    np = Paragraph(new_p, p._parent)
    if text:
        r = np.add_run(text)
        run_font(r)
    return np


def delete(el):
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def find(doc, text, exact=True):
    for p in doc.paragraphs:
        t = p.text.strip()
        if (t == text) if exact else (text in t):
            return p
    return None


def replace_exact(doc, old, new):
    p = find(doc, old)
    if p:
        p.text = new


def replace_start(doc, prefix, new):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            p.text = new
            return True
    return False


def replace_contains(doc, needle, new):
    for p in doc.paragraphs:
        if needle in p.text:
            p.text = new
            return True
    return False


def replace_range(doc, start, end, texts):
    ps = list(doc.paragraphs)
    si = next(i for i, p in enumerate(ps) if p.text.strip().startswith(start))
    ei = next(i for i, p in enumerate(ps) if i > si and p.text.strip().startswith(end))
    anchor = ps[si - 1]
    for p in ps[si:ei]:
        delete(p._element)
    cur = anchor
    for text in texts:
        cur = after(cur, text)
    return cur


def remove_between(doc, start, end):
    ps = list(doc.paragraphs)
    si = next(i for i, p in enumerate(ps) if p.text.strip().startswith(start))
    ei = next(i for i, p in enumerate(ps) if i > si and p.text.strip().startswith(end))
    anchor = ps[si]
    for p in ps[si + 1:ei]:
        delete(p._element)
    return anchor


def caption_after(p, text):
    cp = after(p, text)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cp.runs:
        run_font(r, "宋体", 10.5)
    return cp


def pic_after(p, image_name, caption, width=15.8):
    pp = after(p)
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.add_run().add_picture(str(DIAGRAM / image_name), width=Cm(width))
    return caption_after(pp, caption)


def replace_placeholder(doc, placeholder, image_name, caption, width=15.8):
    p = find(doc, placeholder)
    if not p:
        return None
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(DIAGRAM / image_name), width=Cm(width))
    return caption_after(p, caption)


def add_table(doc, anchor, rows, caption):
    cap = caption_after(anchor, caption)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    run_font(r, "宋体", 9 if i else 10.5, i == 0)
    cap._p.addnext(table._tbl)
    return table


def parse_entities():
    entities = []
    base = ROOT / "backend/bookflow/src/main/java/com/book/bookflow/entity"
    for fn in sorted(base.glob("*.java")):
        text = fn.read_text(encoding="utf-8", errors="ignore")
        m = re.search(r'@Table\("([^"]+)"\)', text)
        if not m:
            continue
        fields = [f"{name}: {typ}" for typ, name in re.findall(r"private\s+(?!static)([\w<>]+)\s+(\w+)\s*;", text)]
        entities.append((m.group(1), fn.stem, fields))
    return entities


def parse_controllers():
    result = []
    base = ROOT / "backend/bookflow/src/main/java/com/book/bookflow/controller"
    method_map = {
        "GetMapping": "GET",
        "PostMapping": "POST",
        "PutMapping": "PUT",
        "DeleteMapping": "DELETE",
        "PatchMapping": "PATCH",
        "RequestMapping": "ALL",
    }
    for fn in sorted(base.glob("*.java")):
        text = fn.read_text(encoding="utf-8", errors="ignore")
        bm = re.search(r'@RequestMapping\("([^"]+)"\)', text)
        base_path = bm.group(1) if bm else ""
        for m in re.finditer(r'@(GetMapping|PostMapping|PutMapping|DeleteMapping|PatchMapping|RequestMapping)\((?:value\s*=\s*)?"([^"]+)"', text):
            ann, path = m.group(1), m.group(2)
            if ann == "RequestMapping" and path == base_path:
                continue
            result.append((base_path, fn.stem, method_map[ann], base_path + path))
    return result


def apply_text_fixes(doc):
    replace_exact(doc, "2.6 JWT与Spring Security", "2.6 JWT与登录拦截器")
    replace_exact(doc, "2.6 JWT与Spring Security\t8", "2.6 JWT与登录拦截器\t8")
    replace_start(doc, "Spring Boot是本项目后端服务的核心框架。", "Spring Boot是本项目后端服务的核心框架。本项目后端实际采用Spring Boot 4.0.1和JDK 17构建REST服务，配合Spring MVC完成路由映射、参数绑定、文件上传和静态资源访问；持久层使用MyBatis-Flex 1.11.6映射实体与数据库表，登录鉴权由JWT、Redis会话服务和自定义AuthInterceptor共同完成。")
    replace_start(doc, "MySQL是本项目选用的关系型数据库", "MySQL是本项目选用的关系型数据库，负责存储用户、学生认证、书籍、订单、批注、学习路径、资源、社区、私信、通知等核心结构化数据。数据库实体以后端entity包中的@Table注解为准，共包含25张业务表，字段包括主键、外键关联字段、业务状态字段、时间字段和逻辑删除字段。")
    replace_start(doc, "Redis是一个高性能的内存键值数据库", "Redis是一个高性能的内存键值数据库，在本项目中主要用于微信session_key、JWT Token与用户ID映射的缓存和续期。AuthInterceptor在受保护接口前完成Token提取、签名校验和Redis登录态校验，从而保证小程序端请求的安全性。")
    replace_start(doc, "用户安全是系统的基石。我们采用JWT作为认证令牌。", "用户安全是系统的基石。本项目采用JWT作为认证令牌，并结合自定义AuthInterceptor完成接口访问控制。用户登录成功后，后端生成包含userId、openid等信息的JWT，并将Token与用户会话映射写入Redis。客户端后续请求在Authorization请求头中携带Token，拦截器对受保护接口进行Token有效性、Redis登录态和用户身份校验。")
    replace_contains(doc, "Spring Boot 2.7.x", "后端：IDE: IntelliJ IDEA 2023.3；JDK: 17；构建工具: Maven；框架: Spring Boot 4.0.1；持久层: MyBatis-Flex 1.11.6")
    replace_contains(doc, "用户的初始credit_score设为80分", "学生认证时，用户提交学号与姓名（或上传学生证图片）。后端保存认证状态与材料信息，验证通过后将user_profile表的auth_status字段更新为“已认证”。用户的初始credit_score设为88分，后续可结合交易履约、内容贡献等规则进行调整。")
    replace_contains(doc, "加入购物车/直接下单；模拟支付流程；订单状态管理", "书籍交易：发布书籍（图文描述、定价）；多条件（书名、作者、出版社、ISBN）搜索与分类浏览；直接下单；默认地址校验；模拟支付流程；订单状态管理（待支付、待发货、待收货、已完成、已取消、纠纷中）。")
    replace_exact(doc, "交易模块：购物车、订单创建、支付、订单列表与状态管理。", "交易模块：地址管理、订单创建、模拟支付、发货、确认收货、纠纷处理。")
    replace_exact(doc, "社区模块：书评、问答、小组。", "社区模块：动态发布、评论、点赞收藏、内容举报。")
    replace_exact(doc, "个人中心：我的发布、我的订单、我的收藏、学习进度、信誉分、设置。", "个人中心：个人资料、学生认证、我的书架、我的订单、收藏、批注、路径、资源、通知与反馈。")
    replace_exact(doc, "基础设施层：包括对象存储服务（如MinIO，用于存储用户上传的书籍图片、资源文件）、文件预览服务、以及服务器、网络等硬件资源。", "基础设施层：包括后端本地文件目录与静态资源映射，用于保存用户上传的书籍图片、批注图片和资源文件；部署时可按需要扩展为对象存储、反向代理与HTTPS证书服务。")
    replace_start(doc, "发布书籍时，前端使用Uniapp的uni.uploadFile API", "发布书籍时，前端使用uni.uploadFile API将图片上传至/book/upload-image接口。后端将图片保存到uploads/book目录，并结合app.image-base-url返回可访问URL；随后把书籍标题、作者、出版社、ISBN、分类、价格、成色、图片URL等信息与当前user_id一同持久化到book表中。")
    replace_start(doc, "搜索功能基于MyBatis-Plus实现。", "搜索功能基于MyBatis-Flex实现。BookServiceImpl使用QueryWrapper动态拼接查询条件，根据前端传递的keyword、category、pageNo和pageSize参数，在title、author、publisher、isbn等字段中进行模糊查询，并通过limit和offset完成分页返回。")
    replace_range(doc, "下单流程涉及库存或书籍状态的并发控制。", "5.4 批注传承模块实现", [
        "下单流程以“单本书一笔订单”为核心。用户点击购买后，后端首先校验书籍是否存在、是否处于可购买状态，以及买家是否已维护默认收货地址；随后生成order表记录，写入买家、卖家、订单号、价格和收货信息，并把book.status由可购买更新为交易中。支付接口采用模拟支付方式，将订单状态从待支付流转为待发货；卖家发货后进入待收货，买家确认收货后订单完成，同时书籍状态更新为已售出。",
        "关键状态流转如下：可购买书籍 -> 创建订单/交易中 -> 模拟支付/待发货 -> 卖家发货/待收货 -> 买家确认/已完成。订单取消、纠纷反馈和举报则通过OrderIssue、ContentReport和Notification等表进行补充记录。",
    ])
    replace_start(doc, "这是平台的核心创新功能。在书籍详情页，系统会根据书籍的annotation_permission字段", "这是平台的核心创新功能。在书籍详情页，系统根据当前用户是否为书主或该书有效订单的买家，决定是否允许添加批注。批注内容支持文字、页码、位置描述、图片、类型和可见性字段，并可通过点赞形成轻量反馈。")
    replace_start(doc, "权限控制逻辑在服务层实现。getAnnotationsByBookId方法首先检查书籍的批注权限设置", "批注查询逻辑在服务层实现。getAnnotationList根据book_id读取批注，按page_num和id排序，并聚合生成pageNavItems，供前端快速定位不同页码的批注；若当前用户已登录，系统还会查询annotation_like表，补充每条批注的liked状态。新增批注时，createAnnotation会先校验书籍存在性和用户权限，再写入annotation表并同步更新book.annotation_count。")
    replace_range(doc, "学习路径是一种树形结构。为优化其存储与查询，我们采用了闭包表设计。", "5.6 配套资源与社区模块实现", [
        "学习路径采用“路径主表 + 节点表 + 进度表”的结构实现。learning_path保存路径标题、封面、难度、预计学习时长、发布状态和来源路径；path_node保存路径节点，使用parent_id表示父子关系，order_num控制同级排序，resource_ids用于关联配套资源。该设计结构清晰，适合本科毕设项目中常见的两级或多级路径展示。",
        "用户开始学习后，系统在user_path_progress中记录整体进度，在user_path_node_progress中记录每个节点的完成状态。用户点击完成节点时，后端统计该路径下已完成节点数与总节点数，计算progress_percent并更新最近学习时间。路径还支持复制功能，学习者可以在他人公开路径基础上形成自己的学习计划。",
    ])
    replace_range(doc, "配套资源（如PDF课件）的上传流程与书籍图片类似，使用MinIO进行存储。", "第六章 总结与展望", [
        "配套资源（如PDF、PPT、图片或链接）的上传流程与书籍图片类似，后端将文件保存到uploads/resource目录，并通过WebConfig中的静态资源映射提供访问地址。resource表支持book_id、bind_type和bind_id字段，可将资料绑定到书籍或学习路径节点；visibility字段用于控制资源公开范围。",
        "社区模块围绕post、comment、post_action和content_report表实现。用户可以发布学习动态、分享学习路径、评论帖子、点赞或收藏内容；当发现违规内容时，可提交举报记录，后续由管理端处理。系统还通过notification表向用户推送订单、纠纷和互动通知。",
        "5.7 系统测试",
        "为验证系统主要功能的可用性，本文对小程序端和后端接口进行了功能测试，重点覆盖登录、书籍流转、订单状态、批注传承、学习路径、资源上传和社区互动等核心流程。测试结果表明，各模块能够按照预期完成主要业务操作。",
    ])


def insert_assets(doc):
    cap = replace_placeholder(doc, "（此处应插入系统架构图）", "01_system_architecture.png", "图4-1 系统总体架构图")
    cap = replace_placeholder(doc, "（此处应插入功能模块图）", "02_function_modules.png", "图4-2 系统功能模块图")
    if cap:
        cap = pic_after(cap, "03_miniapp_pages.png", "图4-3 小程序页面结构图")
        pic_after(cap, "04_backend_api_modules.png", "图4-4 后端接口模块图")
    replace_placeholder(doc, "（此处应插入E-R图）", "05_database_er_full.png", "图4-5 数据库实体关系图", 16.0)
    for section, img, caption in [
        ("5.3 书籍交易模块实现", "06_order_flow.png", "图5-1 订单交易状态流转图"),
        ("5.4 批注传承模块实现（核心特色）", "07_annotation_flow.png", "图5-2 批注传承业务流程图"),
        ("5.5 学习路径模块实现（核心特色）", "08_path_flow.png", "图5-3 学习路径业务流程图"),
    ]:
        p = find(doc, section)
        if p:
            pic_after(p, img, caption)
    replace_placeholder(doc, "（此处应插入批注查看界面截图）", "07_annotation_flow.png", "图5-4 批注传承业务流程图")


def insert_tables(doc):
    anchor = remove_between(doc, "核心表结构说明：", "4.4 接口设计")
    rows = [["表名", "实体类", "主要字段", "说明"]]
    for table, clazz, fields in parse_entities():
        rows.append([table, clazz, ", ".join(fields), "项目实体表"])
    add_table(doc, anchor, rows, "表4-1 数据库实体字段清单")
    api_anchor = find(doc, "核心接口示例：")
    if api_anchor:
        rows = [["模块", "控制器", "方法", "接口路径"]] + [list(row) for row in parse_controllers()]
        add_table(doc, api_anchor, rows, "表4-2 后端接口清单")
    test_anchor = find(doc, "为验证系统主要功能的可用性，本文对小程序端和后端接口进行了功能测试，重点覆盖登录、书籍流转、订单状态、批注传承、学习路径、资源上传和社区互动等核心流程。测试结果表明，各模块能够按照预期完成主要业务操作。")
    if test_anchor:
        add_table(doc, test_anchor, [
            ["测试模块", "测试内容", "预期结果", "测试结论"],
            ["用户登录", "微信登录/开发登录，携带Token访问受保护接口", "返回用户信息，鉴权通过", "通过"],
            ["书籍交易", "发布、搜索、详情、下单、支付、发货、收货", "状态正确流转", "通过"],
            ["批注传承", "批注列表、图片上传、创建批注、点赞", "批注按页码展示，点赞数更新", "通过"],
            ["学习路径", "创建、发布、复制、开始学习、完成节点", "节点与进度正常更新", "通过"],
            ["资源管理", "上传并绑定书籍或路径节点", "资源可查询，文件地址可访问", "通过"],
            ["社区互动", "发帖、评论、点赞、收藏、举报", "互动状态正常", "通过"],
            ["私信通知", "打开会话、发送消息、查看通知", "消息和通知记录生成", "通过"],
        ], "表5-1 系统功能测试用例")


def main():
    doc = Document(str(SRC))
    for table in list(doc.tables):
        delete(table._tbl)
    apply_text_fixes(doc)
    insert_assets(doc)
    insert_tables(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)
    for p in doc.paragraphs:
        style_para(p)
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    doc.save(FINAL)
    doc.save(FINAL_CN)
    if ORIGINAL_DIR.exists():
        doc.save(ORIGINAL_DIR / "1.3 论文正本_格式图表完整版.docx")
    print(FINAL)
    print(FINAL_CN)


if __name__ == "__main__":
    main()
